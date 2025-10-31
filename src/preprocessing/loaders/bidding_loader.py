"""
Bidding Document Loader for Vietnamese Bidding Documents
Handles: Hồ sơ mời thầu (HSMT), E-HSDT, etc.

Refactored from: archive/preprocessing_v1/bidding_preprocessing/extractors/bidding_extractor.py
Integrated with V2 Unified Schema
"""

from __future__ import annotations
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, TYPE_CHECKING
from dataclasses import dataclass

if TYPE_CHECKING:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class RawBiddingContent:
    """
    Raw extracted content from bidding documents.
    Pre-schema data structure (similar to RawDocxContent).
    """

    text: str
    metadata: Dict
    tables: List[Dict]
    structure: List[Dict]  # Sections, chapters, forms
    statistics: Dict


class BiddingLoader:
    """
    Loader for Vietnamese bidding documents (Hồ sơ mời thầu).

    Supports bidding document types:
    - Construction (Xây lắp)
    - Goods (Hàng hóa)
    - Consulting services (Tư vấn)
    - Non-consulting services (Phi tư vấn)
    - EPC (Engineering, Procurement, Construction)
    - EP, EC, PC variants
    - Equipment lease (Máy đặt máy mượn)
    - Online bidding (Chào giá trực tuyến)
    - Online procurement (Mua sắm trực tuyến)

    Example:
        loader = BiddingLoader()
        content = loader.load("path/to/HSMT.docx")
        print(f"Bidding type: {content.metadata['bidding_type']}")
    """

    def __init__(self, preserve_formatting: bool = True):
        """
        Args:
            preserve_formatting: Whether to preserve text formatting
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )

        self.preserve_formatting = preserve_formatting

        # Bidding document type detection patterns
        self.bidding_type_patterns = {
            "construction": re.compile(r"(?i)(xây\s*lắp|xây\s*dựng|thi\s*công)"),
            "goods": re.compile(r"(?i)(hàng\s*hóa|thiết\s*bị|vật\s*tư)"),
            "consulting": re.compile(r"(?i)(tư\s*vấn|dịch\s*vụ\s*tư\s*vấn)"),
            "non_consulting": re.compile(
                r"(?i)(phi\s*tư\s*vấn|dịch\s*vụ\s*phi\s*tư\s*vấn)"
            ),
            "epc": re.compile(r"(?i)(epc|engineering.*procurement.*construction)"),
            "ep": re.compile(r"(?i)(\bep\b|engineering.*procurement)"),
            "ec": re.compile(r"(?i)(\bec\b|engineering.*construction)"),
            "pc": re.compile(r"(?i)(\bpc\b|procurement.*construction)"),
            "equipment_lease": re.compile(
                r"(?i)(máy\s*đặt|máy\s*mượn|thuê\s*thiết\s*bị)"
            ),
            "online_bidding": re.compile(
                r"(?i)(chào\s*giá\s*trực\s*tuyến|đấu\s*thầu\s*trực\s*tuyến)"
            ),
            "online_procurement": re.compile(r"(?i)(mua\s*sắm\s*trực\s*tuyến)"),
        }

        # Structure patterns (similar to legal docs but for bidding)
        self.structure_patterns = {
            "section": re.compile(
                r"^(PHẦN|Phần)\s+([IVXLCDM]+|\d+)[\.\:]?\s*(.+)$", re.IGNORECASE
            ),
            "chapter": re.compile(
                r"^(CHƯƠNG|Chương)\s+([IVXLCDM]+|\d+)[\.\:]?\s*(.+)$", re.IGNORECASE
            ),
            "article": re.compile(r"^(Điều|ĐIỀU)\s+(\d+[a-z]?)[\.\:]?\s*(.+)$"),
            "clause": re.compile(r"^(\d+)[\.\)]?\s+(.+)"),
            "form": re.compile(r"^(Mẫu|MẪU)\s+(\d+[A-Z]?)[\.\:]?\s*(.+)$"),
            "appendix": re.compile(
                r"^(Phụ\s*lục|PHỤ\s*LỤC)\s+(\d+[A-Z]?)[\.\:]?\s*(.+)$", re.IGNORECASE
            ),
            "table_title": re.compile(r"^(Bảng|BẢNG)\s+(\d+)[\.\:]?\s*(.+)$"),
        }

        # Bidding terminology for detection
        self.bidding_keywords = {
            "invitation": [
                "mời thầu",
                "mời chào giá",
                "mời quan tâm",
                "hsmt",
                "e-hsdt",
            ],
            "contractor": ["nhà thầu", "đơn vị thầu", "thầu phụ"],
            "owner": ["chủ đầu tư", "bên mời thầu", "chủ dự án"],
            "contract": ["hợp đồng", "gói thầu", "dự án"],
            "evaluation": ["đánh giá", "xét thầu", "chấm thầu"],
            "technical": ["kỹ thuật", "chuyên môn", "công nghệ", "yêu cầu kỹ thuật"],
            "financial": ["tài chính", "kinh tế", "chi phí", "giá dự thầu"],
            "qualification": ["năng lực", "kinh nghiệm", "trình độ", "đủ điều kiện"],
        }

    def load(self, file_path: str) -> RawBiddingContent:
        """
        Load and extract bidding document.

        Args:
            file_path: Path to bidding document (.docx)

        Returns:
            RawBiddingContent with extracted text, metadata, tables, structure
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.suffix.lower() != ".docx":
            raise ValueError(
                f"Only .docx files supported. Got: {file_path.suffix}\n"
                f"For .doc files, convert to .docx first."
            )

        return self._extract_docx(file_path)

    def _extract_docx(self, file_path: Path) -> RawBiddingContent:
        """Extract content from DOCX bidding document"""
        doc = Document(str(file_path))

        # Extract text and structure
        text, structure = self._extract_text_and_structure(doc)

        # Extract tables
        tables = self._extract_tables(doc)

        # Extract metadata
        metadata = self._extract_metadata(doc, file_path, text)

        # Calculate statistics
        statistics = self._calculate_statistics(text, structure, tables, metadata)

        return RawBiddingContent(
            text=text,
            metadata=metadata,
            tables=tables,
            structure=structure,
            statistics=statistics,
        )

    def _extract_text_and_structure(self, doc: Document) -> Tuple[str, List[Dict]]:
        """
        Extract text and detect bidding document structure.

        Returns:
            (full_text, structure_list)
        """
        text_parts = []
        structure = []

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            # Check for structure elements
            structure_item = self._detect_structure_element(para_text)
            if structure_item:
                structure.append(structure_item)

            text_parts.append(para_text)

        full_text = "\n\n".join(text_parts)
        return full_text, structure

    def _detect_structure_element(self, text: str) -> Optional[Dict]:
        """Detect if text is a structure element (section, form, etc.)"""
        for element_type, pattern in self.structure_patterns.items():
            match = pattern.match(text)
            if match:
                if element_type in [
                    "section",
                    "chapter",
                    "form",
                    "appendix",
                    "table_title",
                ]:
                    return {
                        "type": element_type,
                        "number": match.group(2) if len(match.groups()) >= 2 else "",
                        "text": match.group(3) if len(match.groups()) >= 3 else text,
                    }
                elif element_type == "article":
                    return {
                        "type": "article",
                        "number": match.group(2),
                        "text": match.group(3) if len(match.groups()) >= 3 else "",
                    }
                elif element_type == "clause":
                    return {
                        "type": "clause",
                        "number": match.group(1),
                        "text": match.group(2),
                    }

        return None

    def _extract_tables(self, doc: Document) -> List[Dict]:
        """Extract all tables from document"""
        tables = []

        for i, table in enumerate(doc.tables):
            table_data = {
                "id": i,
                "title": f"Bảng {i+1}",
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "content": [],
            }

            # Extract table content (first 10 rows for preview)
            for row_idx, row in enumerate(table.rows[:10]):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace("\n", " ")
                    row_data.append(cell_text)
                table_data["content"].append(row_data)

                # Try to detect table title from first row
                if row_idx == 0 and row_data:
                    first_cell = row_data[0]
                    # Check if first cell looks like a title
                    if any(
                        keyword in first_cell.lower()
                        for keyword in ["bảng", "mẫu", "phụ lục"]
                    ):
                        table_data["title"] = first_cell

            tables.append(table_data)

        return tables

    def _extract_metadata(self, doc: Document, file_path: Path, text: str) -> Dict:
        """Extract bidding document metadata"""
        core_props = doc.core_properties

        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path.absolute()),
            "title": core_props.title or file_path.stem,
            "author": core_props.author or "",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": (
                core_props.modified.isoformat() if core_props.modified else None
            ),
        }

        # Detect bidding type
        bidding_type = self._detect_bidding_type(text)
        metadata["bidding_type"] = bidding_type

        # Detect bidding package info
        package_info = self._detect_package_info(text)
        metadata.update(package_info)

        # Detect if contains specific sections
        metadata["contains_forms"] = self._contains_forms(text)
        metadata["contains_technical_specs"] = self._contains_technical_specs(text)
        metadata["contains_financial_reqs"] = self._contains_financial_requirements(
            text
        )
        metadata["contains_evaluation_criteria"] = self._contains_evaluation_criteria(
            text
        )

        return metadata

    def _detect_bidding_type(self, text: str) -> str:
        """Detect type of bidding document"""
        text_lower = text.lower()

        # Check each pattern
        for bid_type, pattern in self.bidding_type_patterns.items():
            if pattern.search(text_lower):
                return bid_type

        return "general"

    def _detect_package_info(self, text: str) -> Dict:
        """Detect bidding package information"""
        info = {
            "package_name": "",
            "package_code": "",
            "owner": "",
        }

        # Try to find package name
        name_patterns = [
            r"(?i)tên\s+gói\s+thầu[:\s]+(.+?)(?:\n|$)",
            r"(?i)gói\s+thầu[:\s]+(.+?)(?:\n|$)",
        ]
        for pattern in name_patterns:
            match = re.search(pattern, text)
            if match:
                info["package_name"] = match.group(1).strip()
                break

        # Try to find package code
        code_patterns = [
            r"(?i)mã\s+gói\s+thầu[:\s]+([A-Z0-9\-/]+)",
            r"(?i)số\s+gói\s+thầu[:\s]+([A-Z0-9\-/]+)",
        ]
        for pattern in code_patterns:
            match = re.search(pattern, text)
            if match:
                info["package_code"] = match.group(1).strip()
                break

        # Try to find owner/investor
        owner_patterns = [
            r"(?i)chủ\s+đầu\s+tư[:\s]+(.+?)(?:\n|$)",
            r"(?i)bên\s+mời\s+thầu[:\s]+(.+?)(?:\n|$)",
        ]
        for pattern in owner_patterns:
            match = re.search(pattern, text)
            if match:
                info["owner"] = match.group(1).strip()
                break

        return info

    def _contains_forms(self, text: str) -> bool:
        """Check if document contains bidding forms"""
        return bool(re.search(r"(?i)(mẫu\s+\d+|biểu\s+mẫu|form)", text))

    def _contains_technical_specs(self, text: str) -> bool:
        """Check if document contains technical specifications"""
        keywords = [
            "yêu cầu kỹ thuật",
            "thông số kỹ thuật",
            "đặc tính kỹ thuật",
            "tiêu chuẩn kỹ thuật",
        ]
        return any(keyword in text.lower() for keyword in keywords)

    def _contains_financial_requirements(self, text: str) -> bool:
        """Check if document contains financial requirements"""
        keywords = [
            "năng lực tài chính",
            "điều kiện tài chính",
            "báo cáo tài chính",
            "tình hình tài chính",
        ]
        return any(keyword in text.lower() for keyword in keywords)

    def _contains_evaluation_criteria(self, text: str) -> bool:
        """Check if document contains evaluation criteria"""
        keywords = [
            "tiêu chuẩn đánh giá",
            "tiêu chí đánh giá",
            "phương pháp đánh giá",
            "xét thầu",
        ]
        return any(keyword in text.lower() for keyword in keywords)

    def _calculate_statistics(
        self, text: str, structure: List[Dict], tables: List[Dict], metadata: Dict
    ) -> Dict:
        """Calculate document statistics"""

        # Count structure elements by type
        structure_counts = {}
        for item in structure:
            item_type = item["type"]
            structure_counts[item_type] = structure_counts.get(item_type, 0) + 1

        return {
            "char_count": len(text),
            "word_count": len(text.split()),
            "line_count": text.count("\n") + 1,
            "table_count": len(tables),
            "section_count": structure_counts.get("section", 0),
            "chapter_count": structure_counts.get("chapter", 0),
            "form_count": structure_counts.get("form", 0),
            "appendix_count": structure_counts.get("appendix", 0),
            "bidding_type": metadata.get("bidding_type", "unknown"),
        }
