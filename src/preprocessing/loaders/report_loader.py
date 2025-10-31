"""
Report Template Loader for Vietnamese Report Documents
Handles: Mẫu báo cáo đánh giá, Báo cáo thẩm định, etc.

Similar to BiddingLoader but for report templates
Integrated with V2 Unified Schema
"""

from __future__ import annotations
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, TYPE_CHECKING
from dataclasses import dataclass

if TYPE_CHECKING:
    from docx import Document
    from docx.table import Table

try:
    from docx import Document
    from docx.table import Table

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class RawReportContent:
    """
    Raw extracted content from report documents.
    Pre-schema data structure.
    """

    text: str
    metadata: Dict
    tables: List[Dict]
    structure: List[Dict]
    statistics: Dict


class ReportLoader:
    """
    Loader for Vietnamese report templates (Mẫu báo cáo).

    Supports report types:
    - Evaluation reports (Báo cáo đánh giá - BCĐG)
    - Appraisal reports (Báo cáo thẩm định - BCTĐ)
    - Technical evaluation reports
    - Financial evaluation reports
    - Qualification evaluation reports

    Example:
        loader = ReportLoader()
        content = loader.load("path/to/bao-cao.docx")
        print(f"Report type: {content.metadata['report_type']}")
    """

    def __init__(self, preserve_formatting: bool = True):
        """
        Args:
            preserve_formatting: Whether to preserve text formatting
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )

        self.preserve_formatting = preserve_formatting

        # Report type detection patterns
        self.report_type_patterns = {
            "evaluation": re.compile(
                r"(?i)(báo\s*cáo\s*đánh\s*giá|bcđg|đánh\s*giá\s*hồ\s*sơ)"
            ),
            "appraisal": re.compile(r"(?i)(báo\s*cáo\s*thẩm\s*định|bctđ|thẩm\s*định)"),
            "technical": re.compile(
                r"(?i)(kỹ\s*thuật|chuyên\s*môn|đánh\s*giá.*kỹ\s*thuật)"
            ),
            "financial": re.compile(
                r"(?i)(tài\s*chính|kinh\s*tế|đánh\s*giá.*tài\s*chính)"
            ),
            "qualification": re.compile(
                r"(?i)(năng\s*lực|kinh\s*nghiệm|đủ\s*điều\s*kiện)"
            ),
        }

        # Structure patterns (similar to bidding docs)
        self.structure_patterns = {
            "section": re.compile(
                r"^(PHẦN|Phần)\s+([IVXLCDM]+|\d+)[\.\:]?\s*(.+)$", re.IGNORECASE
            ),
            "chapter": re.compile(
                r"^(CHƯƠNG|Chương|MỤC|Mục)\s+([IVXLCDM]+|\d+)[\.\:]?\s*(.+)$",
                re.IGNORECASE,
            ),
            "article": re.compile(r"^(Điều|ĐIỀU)\s+(\d+[a-z]?)[\.\:]?\s*(.+)$"),
            "clause": re.compile(r"^(\d+)[\.\)]?\s+(.+)"),
            "table_title": re.compile(r"^(Bảng|BẢNG|Biểu)\s+(\d+)[\.\:]?\s*(.+)$"),
        }

    def load(self, file_path: str) -> RawReportContent:
        """
        Load and extract report document.

        Args:
            file_path: Path to report document (.docx)

        Returns:
            RawReportContent with extracted text, metadata, tables, structure
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.suffix.lower() not in [".docx", ".doc"]:
            raise ValueError(
                f"Only .docx/.doc files supported. Got: {file_path.suffix}"
            )

        return self._extract_docx(file_path)

    def _extract_docx(self, file_path: Path) -> RawReportContent:
        """Extract content from DOCX report document"""
        doc = Document(str(file_path))

        # Extract text and structure
        text, structure = self._extract_text_and_structure(doc)

        # Extract tables (reports usually have many tables)
        tables = self._extract_tables(doc)

        # Extract metadata
        metadata = self._extract_metadata(doc, file_path, text)

        # Calculate statistics
        statistics = self._calculate_statistics(text, structure, tables, metadata)

        return RawReportContent(
            text=text,
            metadata=metadata,
            tables=tables,
            structure=structure,
            statistics=statistics,
        )

    def _extract_text_and_structure(self, doc: Document) -> Tuple[str, List[Dict]]:
        """Extract text and detect report structure"""
        text_parts = []
        structure = []

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            # Check for structure elements
            structure_item = self._detect_structure_element(para_text)
            if structure_item:
                structure.append(structure_item)

            text_parts.append(para_text)

        full_text = "\n\n".join(text_parts)
        return full_text, structure

    def _detect_structure_element(self, text: str) -> Optional[Dict]:
        """Detect if text is a structure element"""
        for element_type, pattern in self.structure_patterns.items():
            match = pattern.match(text)
            if match:
                if element_type in ["section", "chapter", "table_title"]:
                    return {
                        "type": element_type,
                        "number": match.group(2) if len(match.groups()) >= 2 else "",
                        "text": match.group(3) if len(match.groups()) >= 3 else text,
                    }
                elif element_type == "article":
                    return {
                        "type": "article",
                        "number": match.group(2),
                        "text": match.group(3) if len(match.groups()) >= 3 else "",
                    }
                elif element_type == "clause":
                    return {
                        "type": "clause",
                        "number": match.group(1),
                        "text": match.group(2),
                    }

        return None

    def _extract_tables(self, doc: Document) -> List[Dict]:
        """Extract all tables from document"""
        tables = []

        for i, table in enumerate(doc.tables):
            table_data = {
                "id": i,
                "title": f"Bảng {i+1}",
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "content": [],
            }

            # Extract table content (first 10 rows)
            for row_idx, row in enumerate(table.rows[:10]):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace("\n", " ")
                    row_data.append(cell_text)
                table_data["content"].append(row_data)

                # Detect table title from first row
                if row_idx == 0 and row_data:
                    first_cell = row_data[0]
                    if any(
                        keyword in first_cell.lower()
                        for keyword in ["bảng", "biểu", "danh sách"]
                    ):
                        table_data["title"] = first_cell

            tables.append(table_data)

        return tables

    def _extract_metadata(self, doc: Document, file_path: Path, text: str) -> Dict:
        """Extract report metadata"""
        core_props = doc.core_properties

        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path.absolute()),
            "title": core_props.title or file_path.stem,
            "author": core_props.author or "",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": (
                core_props.modified.isoformat() if core_props.modified else None
            ),
        }

        # Detect report type
        report_type = self._detect_report_type(text)
        metadata["report_type"] = report_type

        # Detect report info
        report_info = self._detect_report_info(text)
        metadata.update(report_info)

        return metadata

    def _detect_report_type(self, text: str) -> str:
        """Detect type of report"""
        text_lower = text.lower()

        # Check each pattern
        for report_type, pattern in self.report_type_patterns.items():
            if pattern.search(text_lower):
                return report_type

        return "general"

    def _detect_report_info(self, text: str) -> Dict:
        """Detect report information"""
        info = {
            "report_title": "",
            "package_name": "",
            "evaluator": "",
        }

        # Try to find report title (usually in first few lines)
        lines = text.split("\n")[:20]
        for line in lines:
            if "báo cáo" in line.lower() and len(line) > 10:
                info["report_title"] = line.strip()
                break

        # Try to find package name
        name_patterns = [
            r"(?i)gói\s+thầu[:\s]+(.+?)(?:\n|$)",
            r"(?i)tên\s+gói\s+thầu[:\s]+(.+?)(?:\n|$)",
        ]
        for pattern in name_patterns:
            match = re.search(pattern, text)
            if match:
                info["package_name"] = match.group(1).strip()
                break

        # Try to find evaluator/appraiser
        evaluator_patterns = [
            r"(?i)đơn\s+vị\s+đánh\s+giá[:\s]+(.+?)(?:\n|$)",
            r"(?i)đơn\s+vị\s+thẩm\s+định[:\s]+(.+?)(?:\n|$)",
        ]
        for pattern in evaluator_patterns:
            match = re.search(pattern, text)
            if match:
                info["evaluator"] = match.group(1).strip()
                break

        return info

    def _calculate_statistics(
        self, text: str, structure: List[Dict], tables: List[Dict], metadata: Dict
    ) -> Dict:
        """Calculate document statistics"""

        # Count structure elements by type
        structure_counts = {}
        for item in structure:
            item_type = item["type"]
            structure_counts[item_type] = structure_counts.get(item_type, 0) + 1

        return {
            "char_count": len(text),
            "word_count": len(text.split()),
            "line_count": text.count("\n") + 1,
            "table_count": len(tables),
            "section_count": structure_counts.get("section", 0),
            "chapter_count": structure_counts.get("chapter", 0),
            "report_type": metadata.get("report_type", "unknown"),
        }
