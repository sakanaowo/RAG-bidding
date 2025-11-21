#!/usr/bin/env python3
"""
Test #2: Cross-Document Type Batch Testing
Timeline: 19:45 - 20:30 (45 minutes)

Test ChunkFactory auto-selection and consistent output across different document types
"""

import sys
from pathlib import Path

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import time
from typing import Dict, Any, List

from docx import Document as DocxDocument
from src.preprocessing.base.models import ProcessedDocument
from src.preprocessing.chunking.chunk_factory import create_chunker


def load_docx_to_processed_doc(filepath: Path, doc_type: str) -> ProcessedDocument:
    """Load DOCX file into ProcessedDocument"""
    doc = DocxDocument(filepath)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    return ProcessedDocument(
        metadata={
            "filename": filepath.name,
            "document_type": doc_type,
            "source_path": str(filepath),
        },
        content={
            "full_text": text,
        },
    )


def test_batch_processing(test_files: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Process multiple documents in batch"""
    print(f"\n{'='*80}")
    print("🔄 BATCH PROCESSING TEST")
    print(f"{'='*80}")
    print(f"Files to process: {len(test_files)}")

    start_time = time.time()
    results = {
        "total_files": len(test_files),
        "processed": 0,
        "failed": 0,
        "total_chunks": 0,
        "by_type": {},
        "chunker_selection": {},
        "execution_time": 0,
    }

    for i, file_info in enumerate(test_files, 1):
        doc_type = file_info["doc_type"]
        filepath = PROJECT_ROOT / file_info["filepath"]

        print(f"\n[{i}/{len(test_files)}] Processing {doc_type}: {filepath.name}")

        try:
            # Load document
            proc_doc = load_docx_to_processed_doc(filepath, doc_type)

            # Auto-select chunker
            chunker = create_chunker(document_type=doc_type)
            chunker_name = chunker.__class__.__name__

            # Track chunker selection
            if doc_type not in results["chunker_selection"]:
                results["chunker_selection"][doc_type] = set()
            results["chunker_selection"][doc_type].add(chunker_name)

            # Chunk document
            chunks = chunker.chunk(proc_doc)

            # Track results
            if doc_type not in results["by_type"]:
                results["by_type"][doc_type] = {
                    "files": 0,
                    "chunks": 0,
                    "avg_chunk_size": 0,
                }

            results["by_type"][doc_type]["files"] += 1
            results["by_type"][doc_type]["chunks"] += len(chunks)
            results["total_chunks"] += len(chunks)
            results["processed"] += 1

            avg_size = (
                sum(len(c.content) for c in chunks) / len(chunks) if chunks else 0
            )
            results["by_type"][doc_type]["avg_chunk_size"] = avg_size

            print(f"   ✅ Success: {len(chunks)} chunks, {chunker_name}")

        except Exception as e:
            results["failed"] += 1
            print(f"   ❌ Failed: {type(e).__name__}: {str(e)[:50]}")

    results["execution_time"] = time.time() - start_time

    # Convert sets to lists for display
    for doc_type in results["chunker_selection"]:
        results["chunker_selection"][doc_type] = list(
            results["chunker_selection"][doc_type]
        )

    return results


def validate_chunker_consistency(results: Dict[str, Any]) -> Dict[str, Any]:
    """Validate that same document types use same chunkers"""
    print(f"\n{'='*80}")
    print("🔍 CHUNKER CONSISTENCY VALIDATION")
    print(f"{'='*80}")

    validation = {
        "consistent": True,
        "issues": [],
    }

    for doc_type, chunkers in results["chunker_selection"].items():
        if len(chunkers) > 1:
            validation["consistent"] = False
            validation["issues"].append(
                f"{doc_type}: Multiple chunkers used: {chunkers}"
            )
            print(f"   ❌ {doc_type}: Inconsistent - {chunkers}")
        else:
            print(f"   ✅ {doc_type}: Consistent - {chunkers[0]}")

    return validation


def validate_output_format(test_files: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Validate all chunks have consistent output format"""
    print(f"\n{'='*80}")
    print("📋 OUTPUT FORMAT VALIDATION")
    print(f"{'='*80}")

    validation = {
        "consistent": True,
        "issues": [],
        "sample_chunk": None,
    }

    all_chunks = []

    for file_info in test_files[:2]:  # Sample first 2 files
        filepath = PROJECT_ROOT / file_info["filepath"]
        doc_type = file_info["doc_type"]

        try:
            proc_doc = load_docx_to_processed_doc(filepath, doc_type)
            chunker = create_chunker(document_type=doc_type)
            chunks = chunker.chunk(proc_doc)
            all_chunks.extend(chunks[:2])  # First 2 chunks from each file
        except Exception as e:
            validation["issues"].append(f"Failed to process {filepath.name}: {e}")

    if not all_chunks:
        validation["consistent"] = False
        validation["issues"].append("No chunks created to validate")
        return validation

    # Check all chunks have same attributes
    first_chunk = all_chunks[0]
    expected_attrs = set(dir(first_chunk))

    for i, chunk in enumerate(all_chunks[1:], 1):
        chunk_attrs = set(dir(chunk))
        if chunk_attrs != expected_attrs:
            validation["consistent"] = False
            missing = expected_attrs - chunk_attrs
            extra = chunk_attrs - expected_attrs
            validation["issues"].append(f"Chunk {i}: Missing {missing}, Extra {extra}")

    # Store sample chunk structure
    validation["sample_chunk"] = {
        "attributes": [a for a in dir(first_chunk) if not a.startswith("_")],
        "has_content": hasattr(first_chunk, "content"),
        "has_metadata": hasattr(first_chunk, "metadata"),
        "has_chunk_id": hasattr(first_chunk, "chunk_id"),
    }

    print(f"   Chunks validated: {len(all_chunks)}")
    print(f"   Expected attributes: {len(expected_attrs)}")

    if validation["consistent"]:
        print(f"   ✅ All chunks have consistent format")
    else:
        print(f"   ❌ Format inconsistencies found")
        for issue in validation["issues"][:3]:
            print(f"      - {issue}")

    return validation


def main():
    """Run cross-type batch processing tests"""
    print("=" * 80)
    print("🧪 CROSS-DOCUMENT TYPE BATCH INTEGRATION TEST")
    print("=" * 80)
    print(f"Time: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Goal: Test ChunkFactory auto-selection and output consistency")
    print()

    # Define test files (2 per type)
    test_files = [
        # Bidding (2 files)
        {
            "doc_type": "bidding",
            "filepath": Path(
                "data/raw/Ho so moi thau/1. Kế hoạch tổng thể LCNT_kế hoạch LCNT/01A. Mẫu HSYC Xây lắp.docx"
            ),
        },
        {
            "doc_type": "bidding",
            "filepath": Path("data/raw/Ho so moi thau/2. HSMT-Hàng hóa/Phần I.docx"),
        },
        # Law (2 files)
        {
            "doc_type": "law",
            "filepath": Path("data/raw/Luat chinh/Luat so 90 2025-qh15.docx"),
        },
        {
            "doc_type": "law",
            "filepath": Path("data/raw/Luat chinh/Luat Dau thau 2023.docx"),
        },
        # Decree (2 files)
        {
            "doc_type": "decree",
            "filepath": Path(
                "data/raw/Nghi dinh/ND 214 - 4.8.2025 - Thay thế NĐ24-original.docx"
            ),
        },
        {
            "doc_type": "decree",
            "filepath": Path("data/raw/Nghi dinh/24-2024-ND-CP.docx"),
        },
        # Circular (2 files)
        {
            "doc_type": "circular",
            "filepath": Path("data/raw/Thong tu/0. Lời văn thông tư.docx"),
        },
        {
            "doc_type": "circular",
            "filepath": Path("data/raw/Thong tu/TT09-ngay02-09-2024 danh so TCDN.docx"),
        },
    ]

    # Filter existing files
    existing_files = []
    for file_info in test_files:
        filepath = PROJECT_ROOT / file_info["filepath"]
        if filepath.exists():
            existing_files.append(file_info)
        else:
            print(f"⚠️  Skipping: {filepath.name} (not found)")

    print(f"\n✅ Found {len(existing_files)}/{len(test_files)} test files")

    # Test 1: Batch Processing
    batch_results = test_batch_processing(existing_files)

    # Test 2: Chunker Consistency
    consistency_validation = validate_chunker_consistency(batch_results)

    # Test 3: Output Format
    format_validation = validate_output_format(existing_files)

    # Summary Report
    print("\n" + "=" * 80)
    print("📊 TEST SUMMARY")
    print("=" * 80)

    print(f"\n📈 Batch Processing:")
    print(
        f"   Files processed: {batch_results['processed']}/{batch_results['total_files']}"
    )
    print(f"   Total chunks:    {batch_results['total_chunks']}")
    print(f"   Failed:          {batch_results['failed']}")
    print(f"   Execution time:  {batch_results['execution_time']:.2f}s")

    print(f"\n📊 By Document Type:")
    for doc_type, stats in batch_results["by_type"].items():
        print(
            f"   {doc_type:12s}: {stats['files']} files, "
            f"{stats['chunks']:4d} chunks, "
            f"{stats['avg_chunk_size']:.0f} avg chars"
        )

    print(f"\n🔧 Chunker Selection:")
    for doc_type, chunkers in batch_results["chunker_selection"].items():
        print(f"   {doc_type:12s}: {chunkers[0]}")

    print(
        f"\n✅ Consistency: {'PASS' if consistency_validation['consistent'] else 'FAIL'}"
    )
    print(f"✅ Format:      {'PASS' if format_validation['consistent'] else 'FAIL'}")

    # Overall success
    all_passed = (
        batch_results["failed"] == 0
        and consistency_validation["consistent"]
        and format_validation["consistent"]
    )

    print("\n" + "=" * 80)
    if all_passed:
        print("🎉 ALL TESTS PASSED! Batch Processing Working!")
        print("=" * 80)
        return 0
    else:
        print("⚠️  SOME TESTS FAILED - See Details Above")
        print("=" * 80)
        return 1


if __name__ == "__main__":
    sys.exit(main())
