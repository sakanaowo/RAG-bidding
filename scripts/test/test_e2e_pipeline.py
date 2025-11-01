#!/usr/bin/env python3
"""
Test #1: End-to-End Pipeline Testing
Timeline: 19:00 - 19:45 (45 minutes)

Test toàn bộ pipeline từ DOCX file → UnifiedLegalChunk
- Extract → Clean → Parse → Chunk → Map → Convert
"""

import sys
from pathlib import Path

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import time
from typing import Dict, Any, List
from dataclasses import asdict

from docx import Document as DocxDocument
from src.preprocessing.base.models import ProcessedDocument
from src.chunking.chunk_factory import create_chunker


def load_docx_to_processed_doc(filepath: Path, doc_type: str) -> ProcessedDocument:
    """Step 1: Extract - Load DOCX and create ProcessedDocument"""
    doc = DocxDocument(filepath)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    return ProcessedDocument(
        metadata={
            "filename": filepath.name,
            "document_type": doc_type,
            "source_path": str(filepath),
        },
        content={
            "full_text": text,
        },
    )


def validate_chunk_schema(chunk) -> Dict[str, Any]:
    """Validate chunk has all required fields"""
    issues = []

    # Required fields for UniversalChunk
    required_fields = ["content", "chunk_id", "document_id", "document_type"]
    for field in required_fields:
        if not hasattr(chunk, field):
            issues.append(f"Missing field: {field}")

    # Validate field types
    if hasattr(chunk, "content") and not isinstance(chunk.content, str):
        issues.append("content must be str")
    if hasattr(chunk, "chunk_id") and not isinstance(chunk.chunk_id, str):
        issues.append("chunk_id must be str")
    if hasattr(chunk, "document_type") and not isinstance(chunk.document_type, str):
        issues.append("document_type must be str")

    return {
        "valid": len(issues) == 0,
        "issues": issues,
        "chunk_id": getattr(chunk, "chunk_id", "unknown"),
    }


def test_single_pipeline(doc_type: str, filepath: Path) -> Dict[str, Any]:
    """Test complete pipeline for one document"""
    print(f"\n{'='*80}")
    print(f"Testing {doc_type.upper()} Pipeline")
    print(f"File: {filepath.name}")
    print(f"{'='*80}")

    start_time = time.time()
    result = {
        "doc_type": doc_type,
        "filename": filepath.name,
        "success": False,
        "chunks_created": 0,
        "valid_chunks": 0,
        "invalid_chunks": 0,
        "issues": [],
        "execution_time": 0,
    }

    try:
        # Step 1: Extract & Parse
        print("📄 Step 1: Extracting document...")
        proc_doc = load_docx_to_processed_doc(filepath, doc_type)
        print(f"   ✅ Extracted {len(proc_doc.content.get('full_text', ''))} chars")

        # Step 2: Select Chunker
        print("🔧 Step 2: Selecting chunker...")
        chunker = create_chunker(document_type=doc_type)
        print(f"   ✅ Selected: {chunker.__class__.__name__}")

        # Step 3: Chunk
        print("✂️  Step 3: Chunking document...")
        chunks = chunker.chunk(proc_doc)
        result["chunks_created"] = len(chunks)
        print(f"   ✅ Created {len(chunks)} chunks")

        # Step 4: Validate Schema
        print("🔍 Step 4: Validating chunk schemas...")
        validation_results = []
        for i, chunk in enumerate(chunks):
            validation = validate_chunk_schema(chunk)
            validation_results.append(validation)
            if validation["valid"]:
                result["valid_chunks"] += 1
            else:
                result["invalid_chunks"] += 1
                result["issues"].extend(
                    [f"Chunk {i}: {issue}" for issue in validation["issues"]]
                )

        # Step 5: Check Quality
        print("📊 Step 5: Checking chunk quality...")
        in_range = sum(1 for c in chunks if 300 <= len(c.content) <= 1500)
        quality = in_range / len(chunks) * 100 if chunks else 0
        print(f"   📏 In-range: {in_range}/{len(chunks)} ({quality:.1f}%)")

        # Summary
        result["success"] = result["invalid_chunks"] == 0
        result["execution_time"] = time.time() - start_time

        status = "✅ PASS" if result["success"] else "❌ FAIL"
        print(f"\n{status} - {result['doc_type']} pipeline")
        print(f"   Chunks: {result['chunks_created']}")
        print(f"   Valid: {result['valid_chunks']}")
        print(f"   Invalid: {result['invalid_chunks']}")
        print(f"   Time: {result['execution_time']:.2f}s")

        if result["issues"]:
            print(f"\n⚠️  Issues found:")
            for issue in result["issues"][:5]:  # Show first 5
                print(f"   - {issue}")
            if len(result["issues"]) > 5:
                issues_count = len(result["issues"])
                print(f"   ... and {issues_count - 5} more")

    except Exception as e:
        result["success"] = False
        result["issues"].append(f"Pipeline error: {type(e).__name__}: {str(e)}")
        result["execution_time"] = time.time() - start_time
        print(f"\n❌ FAIL - Pipeline crashed: {e}")

    return result


def main():
    """Run E2E pipeline tests"""
    print("=" * 80)
    print("🧪 END-TO-END PIPELINE INTEGRATION TEST")
    print("=" * 80)
    print(f"Time: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Goal: Test complete pipeline for all document types")
    print()

    # Define test cases
    test_cases = [
        {
            "doc_type": "bidding",
            "filepath": Path(
                "data/raw/Ho so moi thau/1. Kế hoạch tổng thể LCNT_kế hoạch LCNT/01A. Mẫu HSYC Xây lắp.docx"
            ),
        },
        {
            "doc_type": "law",
            "filepath": Path("data/raw/Luat chinh/Luat so 90 2025-qh15.docx"),
        },
        {
            "doc_type": "decree",
            "filepath": Path(
                "data/raw/Nghi dinh/ND 214 - 4.8.2025 - Thay thế NĐ24-original.docx"
            ),
        },
        {
            "doc_type": "circular",
            "filepath": Path("data/raw/Thong tu/0. Lời văn thông tư.docx"),
        },
    ]

    # Run tests
    results = []
    for test_case in test_cases:
        filepath = PROJECT_ROOT / test_case["filepath"]
        if not filepath.exists():
            print(f"\n⚠️  Skipping {test_case['doc_type']}: File not found")
            continue

        result = test_single_pipeline(test_case["doc_type"], filepath)
        results.append(result)

    # Summary Report
    print("\n" + "=" * 80)
    print("📊 TEST SUMMARY")
    print("=" * 80)

    total_tests = len(results)
    passed_tests = sum(1 for r in results if r["success"])
    failed_tests = total_tests - passed_tests

    print(f"\nTests Run: {total_tests}")
    print(
        f"Passed:    {passed_tests}/{total_tests} ({passed_tests/total_tests*100:.1f}%)"
    )
    print(f"Failed:    {failed_tests}/{total_tests}")

    print("\n📋 Results by Document Type:")
    for result in results:
        status = "✅" if result["success"] else "❌"
        print(
            f"   {status} {result['doc_type']:12s} | "
            f"Chunks: {result['chunks_created']:3d} | "
            f"Valid: {result['valid_chunks']:3d} | "
            f"Time: {result['execution_time']:.2f}s"
        )

    # Overall success
    print("\n" + "=" * 80)
    if passed_tests == total_tests:
        print("🎉 ALL TESTS PASSED! E2E Pipeline Working!")
        print("=" * 80)
        return 0
    else:
        print(f"⚠️  {failed_tests} TEST(S) FAILED - Need Investigation")
        print("=" * 80)

        # Show issues
        print("\n🔍 Issues Found:")
        for result in results:
            if not result["success"] and result["issues"]:
                print(f"\n{result['doc_type'].upper()}:")
                for issue in result["issues"][:3]:
                    print(f"   - {issue}")

        return 1


if __name__ == "__main__":
    sys.exit(main())
