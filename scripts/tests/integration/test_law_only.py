#!/usr/bin/env python3
"""Test chunking quality for Law documents only."""

import sys
from pathlib import Path

# Add project root to path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from docx import Document
from src.preprocessing.chunking.hierarchical_chunker import HierarchicalChunker
from src.preprocessing.base.models import ProcessedDocument


def test_law_documents():
    """Test law document loading and chunking."""

    law_dir = project_root / "data/raw/Luat chinh"
    law_files = list(law_dir.glob("*.docx"))

    print("=" * 80)
    print("LAW DOCUMENTS TEST")
    print("=" * 80)
    print(f"\n📁 Found {len(law_files)} law files:\n")

    chunker = HierarchicalChunker(
        min_size=300,
        max_size=1500,
        target_size=800,
    )

    total_docs = 0
    loaded_docs = 0
    total_chunks = 0
    in_range_chunks = 0

    for law_file in sorted(law_files):
        total_docs += 1
        print(f"\n📄 {law_file.name}")
        print("-" * 80)

        try:
            # Try to load document
            doc = Document(law_file)
            loaded_docs += 1
            print(f"   ✅ Loaded successfully")
            print(f"   📖 Paragraphs: {len(doc.paragraphs)}")

            # Try to chunk
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

            # Create ProcessedDocument
            proc_doc = ProcessedDocument(
                metadata={
                    "filename": law_file.name,
                    "document_type": "law",
                },
                content={
                    "full_text": text,
                },
            )

            chunks = chunker.chunk(proc_doc)

            # Calculate quality
            chunks_in_range = sum(1 for c in chunks if 300 <= len(c.content) <= 1500)
            quality = chunks_in_range / len(chunks) * 100 if chunks else 0

            total_chunks += len(chunks)
            in_range_chunks += chunks_in_range

            status = "✅" if quality >= 80 else "⚠️"
            print(
                f"   {status} Chunks: {len(chunks)} total, {chunks_in_range} in-range ({quality:.1f}%)"
            )
            print(
                f"   📏 Size range: {min(len(c.content) for c in chunks)}-{max(len(c.content) for c in chunks)} chars"
            )
            print(
                f"   📊 Avg size: {sum(len(c.content) for c in chunks) // len(chunks)} chars"
            )

        except Exception as e:
            print(f"   ❌ ERROR: {type(e).__name__}: {str(e)[:100]}")

    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"Files found: {total_docs}")
    print(
        f"Files loaded: {loaded_docs}/{total_docs} ({loaded_docs/total_docs*100:.1f}%)"
    )
    print(f"Total chunks: {total_chunks}")
    print(
        f"In-range chunks: {in_range_chunks}/{total_chunks} ({in_range_chunks/total_chunks*100:.1f}%)"
    )
    print("=" * 80)

    return loaded_docs == total_docs


if __name__ == "__main__":
    success = test_law_documents()
    sys.exit(0 if success else 1)
