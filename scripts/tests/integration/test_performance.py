#!/usr/bin/env python3
"""
Test #5: Performance & Scalability
Timeline: 22:00 - 22:45 (45 minutes)

Test system performance with larger datasets:
- Processing speed (chunks/second)
- Memory usage tracking
- Batch processing efficiency
- Identify bottlenecks
"""

import sys
from pathlib import Path

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import time
import psutil
import os
from typing import Dict, Any, List
from dataclasses import asdict

from docx import Document as DocxDocument
from src.preprocessing.base.models import ProcessedDocument
from src.preprocessing.chunking.chunk_factory import create_chunker


def get_memory_usage() -> float:
    """Get current process memory usage in MB"""
    process = psutil.Process(os.getpid())
    return process.memory_info().rss / 1024 / 1024  # Convert to MB


def collect_all_test_files() -> List[Dict[str, Any]]:
    """Collect all available test files"""
    test_files = []

    # Bidding documents
    bidding_dir = PROJECT_ROOT / "data/raw/Ho so moi thau"
    for subdir in ["1. Kế hoạch tổng thể LCNT_kế hoạch LCNT", "2. HSMT-Hàng hóa"]:
        dir_path = bidding_dir / subdir
        if dir_path.exists():
            for docx_file in dir_path.glob("*.docx"):
                test_files.append(
                    {
                        "doc_type": "bidding",
                        "filepath": docx_file.relative_to(PROJECT_ROOT),
                    }
                )

    # Law documents
    law_dir = PROJECT_ROOT / "data/raw/Luat chinh"
    if law_dir.exists():
        for docx_file in law_dir.glob("*.docx"):
            test_files.append(
                {
                    "doc_type": "law",
                    "filepath": docx_file.relative_to(PROJECT_ROOT),
                }
            )

    # Decree documents
    decree_dir = PROJECT_ROOT / "data/raw/Nghi dinh"
    if decree_dir.exists():
        for docx_file in decree_dir.glob("*.docx"):
            test_files.append(
                {
                    "doc_type": "decree",
                    "filepath": docx_file.relative_to(PROJECT_ROOT),
                }
            )

    # Circular documents
    circular_dir = PROJECT_ROOT / "data/raw/Thong tu"
    if circular_dir.exists():
        for docx_file in circular_dir.glob("*.docx"):
            test_files.append(
                {
                    "doc_type": "circular",
                    "filepath": docx_file.relative_to(PROJECT_ROOT),
                }
            )

    return test_files


def test_single_document_speed(filepath: Path, doc_type: str) -> Dict[str, Any]:
    """Test processing speed for a single document"""
    result = {
        "filename": filepath.name,
        "doc_type": doc_type,
        "success": False,
        "load_time": 0,
        "chunk_time": 0,
        "total_time": 0,
        "chunks_created": 0,
        "chars_processed": 0,
        "memory_before": 0,
        "memory_after": 0,
        "memory_delta": 0,
    }

    try:
        result["memory_before"] = get_memory_usage()
        start_total = time.time()

        # Step 1: Load document
        start_load = time.time()
        doc = DocxDocument(filepath)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        result["chars_processed"] = len(text)

        proc_doc = ProcessedDocument(
            metadata={
                "filename": filepath.name,
                "document_type": doc_type,
                "source_path": str(filepath),
            },
            content={"full_text": text},
        )
        result["load_time"] = time.time() - start_load

        # Step 2: Chunk document
        start_chunk = time.time()
        chunker = create_chunker(document_type=doc_type)
        chunks = chunker.chunk(proc_doc)
        result["chunk_time"] = time.time() - start_chunk

        result["chunks_created"] = len(chunks)
        result["total_time"] = time.time() - start_total
        result["memory_after"] = get_memory_usage()
        result["memory_delta"] = result["memory_after"] - result["memory_before"]
        result["success"] = True

    except Exception as e:
        result["error"] = f"{type(e).__name__}: {str(e)[:50]}"

    return result


def test_batch_processing(
    test_files: List[Dict[str, Any]], max_files: int = 20
) -> Dict[str, Any]:
    """Test batch processing performance"""
    print(f"\n{'='*80}")
    print(f"🔄 BATCH PROCESSING TEST ({max_files} files)")
    print(f"{'='*80}")

    batch_result = {
        "total_files": 0,
        "processed": 0,
        "failed": 0,
        "total_chunks": 0,
        "total_chars": 0,
        "total_time": 0,
        "memory_start": 0,
        "memory_peak": 0,
        "memory_end": 0,
        "individual_results": [],
    }

    batch_result["memory_start"] = get_memory_usage()
    start_time = time.time()

    # Process files
    for i, file_info in enumerate(test_files[:max_files], 1):
        filepath = PROJECT_ROOT / file_info["filepath"]
        doc_type = file_info["doc_type"]

        print(
            f"[{i}/{min(len(test_files), max_files)}] Processing {doc_type}: {filepath.name[:40]}"
        )

        result = test_single_document_speed(filepath, doc_type)
        batch_result["individual_results"].append(result)

        if result["success"]:
            batch_result["processed"] += 1
            batch_result["total_chunks"] += result["chunks_created"]
            batch_result["total_chars"] += result["chars_processed"]

            # Track peak memory
            current_mem = get_memory_usage()
            if current_mem > batch_result["memory_peak"]:
                batch_result["memory_peak"] = current_mem

            print(
                f"   ✅ {result['chunks_created']:3d} chunks in {result['total_time']:.2f}s "
                f"({result['chars_processed']:,} chars)"
            )
        else:
            batch_result["failed"] += 1
            print(f"   ❌ Failed: {result.get('error', 'Unknown error')}")

    batch_result["total_files"] = min(len(test_files), max_files)
    batch_result["total_time"] = time.time() - start_time
    batch_result["memory_end"] = get_memory_usage()

    return batch_result


def analyze_performance(batch_result: Dict[str, Any]) -> Dict[str, Any]:
    """Analyze performance metrics"""
    print(f"\n{'='*80}")
    print("📊 PERFORMANCE ANALYSIS")
    print(f"{'='*80}")

    analysis = {
        "throughput": {
            "files_per_second": 0,
            "chunks_per_second": 0,
            "chars_per_second": 0,
        },
        "memory": {
            "start_mb": batch_result["memory_start"],
            "peak_mb": batch_result["memory_peak"],
            "end_mb": batch_result["memory_end"],
            "total_increase_mb": 0,
        },
        "timing": {
            "avg_file_time": 0,
            "slowest_file": None,
            "fastest_file": None,
        },
        "bottlenecks": [],
    }

    # Calculate throughput
    if batch_result["total_time"] > 0:
        analysis["throughput"]["files_per_second"] = (
            batch_result["processed"] / batch_result["total_time"]
        )
        analysis["throughput"]["chunks_per_second"] = (
            batch_result["total_chunks"] / batch_result["total_time"]
        )
        analysis["throughput"]["chars_per_second"] = (
            batch_result["total_chars"] / batch_result["total_time"]
        )

    print(f"\n⚡ Throughput:")
    print(f"   Files:  {analysis['throughput']['files_per_second']:.1f} files/sec")
    print(f"   Chunks: {analysis['throughput']['chunks_per_second']:.0f} chunks/sec")
    print(f"   Chars:  {analysis['throughput']['chars_per_second']:,.0f} chars/sec")

    # Memory analysis
    analysis["memory"]["total_increase_mb"] = (
        batch_result["memory_end"] - batch_result["memory_start"]
    )

    print(f"\n💾 Memory Usage:")
    print(f"   Start: {analysis['memory']['start_mb']:.1f} MB")
    print(f"   Peak:  {analysis['memory']['peak_mb']:.1f} MB")
    print(f"   End:   {analysis['memory']['end_mb']:.1f} MB")
    print(f"   Delta: {analysis['memory']['total_increase_mb']:.1f} MB")

    # Timing analysis
    successful_results = [r for r in batch_result["individual_results"] if r["success"]]

    if successful_results:
        analysis["timing"]["avg_file_time"] = sum(
            r["total_time"] for r in successful_results
        ) / len(successful_results)

        slowest = max(successful_results, key=lambda r: r["total_time"])
        fastest = min(successful_results, key=lambda r: r["total_time"])

        analysis["timing"]["slowest_file"] = {
            "name": slowest["filename"],
            "time": slowest["total_time"],
        }
        analysis["timing"]["fastest_file"] = {
            "name": fastest["filename"],
            "time": fastest["total_time"],
        }

        print(f"\n⏱️  Timing:")
        print(f"   Average:  {analysis['timing']['avg_file_time']:.2f}s per file")
        print(f"   Slowest:  {slowest['filename'][:40]} ({slowest['total_time']:.2f}s)")
        print(f"   Fastest:  {fastest['filename'][:40]} ({fastest['total_time']:.2f}s)")

    # Identify bottlenecks
    if analysis["memory"]["peak_mb"] > 500:
        analysis["bottlenecks"].append("High memory usage (>500MB)")

    if analysis["timing"]["avg_file_time"] > 2.0:
        analysis["bottlenecks"].append("Slow processing (>2s per file)")

    if batch_result["total_time"] > 30:
        analysis["bottlenecks"].append(
            f"Long total time ({batch_result['total_time']:.0f}s for {batch_result['processed']} files)"
        )

    if analysis["bottlenecks"]:
        print(f"\n⚠️  Bottlenecks Identified:")
        for bottleneck in analysis["bottlenecks"]:
            print(f"   - {bottleneck}")
    else:
        print(f"\n✅ No significant bottlenecks detected!")

    return analysis


def main():
    """Run performance tests"""
    print("=" * 80)
    print("🧪 PERFORMANCE & SCALABILITY TEST")
    print("=" * 80)
    print(f"Time: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Goal: Benchmark processing speed and memory usage")
    print()

    # Collect test files
    print("📂 Collecting test files...")
    test_files = collect_all_test_files()
    print(f"   Found {len(test_files)} total files")

    if not test_files:
        print("\n❌ No test files found - cannot proceed")
        return 1

    # Run batch processing test
    batch_result = test_batch_processing(test_files, max_files=20)

    # Analyze performance
    analysis = analyze_performance(batch_result)

    # Summary Report
    print("\n" + "=" * 80)
    print("📊 TEST SUMMARY")
    print("=" * 80)

    print(
        f"\nFiles Processed: {batch_result['processed']}/{batch_result['total_files']}"
    )
    print(f"Total Chunks:    {batch_result['total_chunks']:,}")
    print(f"Total Time:      {batch_result['total_time']:.2f}s")
    print(
        f"Success Rate:    {batch_result['processed']/batch_result['total_files']*100:.1f}%"
    )

    # Performance criteria
    print("\n🎯 Performance Criteria:")

    criteria_met = []

    # Criterion 1: Speed
    speed_ok = batch_result["total_time"] < 30
    criteria_met.append(speed_ok)
    print(
        f"   {'✅' if speed_ok else '❌'} Total time < 30s: {batch_result['total_time']:.1f}s"
    )

    # Criterion 2: Memory
    memory_ok = batch_result["memory_peak"] < 500
    criteria_met.append(memory_ok)
    print(
        f"   {'✅' if memory_ok else '❌'} Peak memory < 500MB: {batch_result['memory_peak']:.1f} MB"
    )

    # Criterion 3: Throughput
    throughput_ok = analysis["throughput"]["chunks_per_second"] > 50
    criteria_met.append(throughput_ok)
    print(
        f"   {'✅' if throughput_ok else '❌'} Throughput > 50 chunks/s: {analysis['throughput']['chunks_per_second']:.0f}"
    )

    # Criterion 4: Success rate
    success_ok = batch_result["processed"] == batch_result["total_files"]
    criteria_met.append(success_ok)
    print(
        f"   {'✅' if success_ok else '❌'} All files processed: {batch_result['processed']}/{batch_result['total_files']}"
    )

    # Overall success
    all_criteria_met = all(criteria_met)

    print("\n" + "=" * 80)
    if all_criteria_met:
        print("🎉 ALL CRITERIA MET! System Performance Excellent!")
        print("=" * 80)
        return 0
    else:
        print("⚠️  SOME CRITERIA NOT MET - See Analysis Above")
        print("=" * 80)
        return 1


if __name__ == "__main__":
    sys.exit(main())
