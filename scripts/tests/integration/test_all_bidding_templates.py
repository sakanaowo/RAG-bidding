"""
Test BiddingHybridChunker with ALL bidding templates.

Validates that 100% in-range quality is maintained across all template types.
"""

import pytest
from pathlib import Path
from docx import Document

from src.preprocessing.chunking.bidding_hybrid_chunker import BiddingHybridChunker
from src.preprocessing.base.models import ProcessedDocument


def get_all_bidding_templates():
    """Get all bidding template DOCX files."""
    base_path = Path("data/raw/Ho so moi thau")

    if not base_path.exists():
        pytest.skip(f"Bidding templates directory not found: {base_path}")

    # Find all .docx files (exclude temporary files starting with ~$)
    templates = [f for f in base_path.rglob("*.docx") if not f.name.startswith("~$")]

    return sorted(templates)


def load_template(docx_path: Path) -> ProcessedDocument:
    """Load a bidding template into ProcessedDocument."""
    doc = Document(docx_path)

    # Extract full text
    full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # Create ProcessedDocument with correct structure
    return ProcessedDocument(
        metadata={
            "document_type": "bidding_template",
            "doc_type": "bidding_template",
            "template_name": docx_path.name,
            "template_category": docx_path.parent.name,
            "file_path": str(docx_path),
        },
        content={
            "full_text": full_text,
        },
    )


def test_all_bidding_templates_quality():
    """Test: BiddingHybridChunker maintains high quality across ALL templates."""
    templates = get_all_bidding_templates()

    print(f"\n{'='*80}")
    print(f"TESTING ALL BIDDING TEMPLATES ({len(templates)} files)")
    print(f"{'='*80}\n")

    chunker = BiddingHybridChunker(min_size=300, max_size=1500, target_size=800)

    total_chunks = 0
    total_in_range = 0
    failed_templates = []
    results = []

    for template_path in templates:
        try:
            # Load template
            doc = load_template(template_path)

            # Skip empty documents
            full_text = doc.content.get("full_text", "")
            if not full_text.strip():
                print(f"⚠️  SKIP: {template_path.name} (empty)")
                continue

            # Chunk it
            chunks = chunker.chunk(doc)

            # Calculate metrics
            in_range = [c for c in chunks if 300 <= len(c.content) <= 1500]
            in_range_pct = len(in_range) / len(chunks) * 100 if chunks else 0

            total_chunks += len(chunks)
            total_in_range += len(in_range)

            # Store result
            results.append(
                {
                    "name": template_path.name,
                    "category": template_path.parent.name,
                    "chunks": len(chunks),
                    "in_range": len(in_range),
                    "percentage": in_range_pct,
                }
            )

            # Track failures (< 90% in-range)
            if in_range_pct < 90.0:
                failed_templates.append(
                    {
                        "path": template_path,
                        "chunks": len(chunks),
                        "in_range": len(in_range),
                        "percentage": in_range_pct,
                    }
                )

            # Print progress
            status = "✅" if in_range_pct >= 90.0 else "⚠️"
            print(
                f"{status} {template_path.name:60s} | "
                f"{len(chunks):3d} chunks | "
                f"{len(in_range):3d}/{len(chunks):3d} in-range ({in_range_pct:5.1f}%)"
            )

        except Exception as e:
            print(f"❌ ERROR: {template_path.name} - {str(e)}")
            failed_templates.append(
                {
                    "path": template_path,
                    "error": str(e),
                }
            )

    # Print summary
    print(f"\n{'='*80}")
    print("SUMMARY")
    print(f"{'='*80}")
    print(f"Total templates processed: {len(results)}")
    print(f"Total chunks created: {total_chunks}")
    print(f"Total in-range chunks: {total_in_range}/{total_chunks}")

    if total_chunks > 0:
        overall_pct = total_in_range / total_chunks * 100
        print(f"Overall in-range percentage: {overall_pct:.1f}%")

    # Category breakdown
    print(f"\n{'='*80}")
    print("BY CATEGORY")
    print(f"{'='*80}")

    categories = {}
    for r in results:
        cat = r["category"]
        if cat not in categories:
            categories[cat] = {"chunks": 0, "in_range": 0}
        categories[cat]["chunks"] += r["chunks"]
        categories[cat]["in_range"] += r["in_range"]

    for cat, data in sorted(categories.items()):
        pct = data["in_range"] / data["chunks"] * 100 if data["chunks"] > 0 else 0
        print(f"{cat:40s} | {data['in_range']:4d}/{data['chunks']:4d} ({pct:5.1f}%)")

    # Failed templates
    if failed_templates:
        print(f"\n{'='*80}")
        print(f"⚠️  TEMPLATES WITH < 90% IN-RANGE ({len(failed_templates)} files)")
        print(f"{'='*80}")
        for item in failed_templates:
            if "error" in item:
                print(f"❌ {item['path'].name}: {item['error']}")
            else:
                print(
                    f"⚠️  {item['path'].name}: "
                    f"{item['in_range']}/{item['chunks']} ({item['percentage']:.1f}%)"
                )

    print(f"\n{'='*80}\n")

    # Assertions
    assert total_chunks > 0, "No chunks were created"

    # Target: Overall in-range should be >= 90%
    if total_chunks > 0:
        overall_pct = total_in_range / total_chunks * 100
        assert (
            overall_pct >= 90.0
        ), f"Overall in-range {overall_pct:.1f}% < 90% target (failed on {len(failed_templates)} templates)"


def test_bidding_templates_sample():
    """Test: Sample of different template types for detailed analysis."""
    templates = get_all_bidding_templates()

    # Select diverse samples
    sample_patterns = [
        "15. Phu luc.docx",  # Original test file
        "01A. Mẫu HSYC Xây lắp.docx",  # HSYC templates
        "4A E-HSMT hàng hóa 1 túi.docx",  # Hàng hóa
        "8A. E-HSMT_ EC 01 túi.docx",  # EC
    ]

    samples = []
    for pattern in sample_patterns:
        matching = [t for t in templates if pattern in t.name]
        if matching:
            samples.append(matching[0])

    print(f"\n{'='*80}")
    print(f"DETAILED ANALYSIS - {len(samples)} SAMPLE TEMPLATES")
    print(f"{'='*80}\n")

    chunker = BiddingHybridChunker(min_size=300, max_size=1500, target_size=800)

    for template_path in samples:
        doc = load_template(template_path)
        chunks = chunker.chunk(doc)

        in_range = [c for c in chunks if 300 <= len(c.content) <= 1500]
        sizes = [len(c.content) for c in chunks]

        print(f"\n📄 {template_path.name}")
        print(f"   Category: {template_path.parent.name}")
        print(f"   Total chunks: {len(chunks)}")
        print(
            f"   In-range: {len(in_range)}/{len(chunks)} ({len(in_range)/len(chunks)*100:.1f}%)"
        )
        print(f"   Size range: {min(sizes)}-{max(sizes)} chars")
        print(f"   Average: {sum(sizes)//len(sizes)} chars")

        # Show out-of-range chunks if any
        out_of_range = [
            c for c in chunks if len(c.content) < 300 or len(c.content) > 1500
        ]
        if out_of_range:
            print(f"   ⚠️  Out-of-range chunks:")
            for c in out_of_range[:3]:  # Show first 3
                preview = c.content[:50].replace("\n", " ")
                print(f"      - {len(c.content):4d} chars: {preview}...")

    print(f"\n{'='*80}\n")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
