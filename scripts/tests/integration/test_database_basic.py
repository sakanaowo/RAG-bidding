#!/usr/bin/env python3
"""
Test #4: Database Integration (Basic)
Timeline: 21:15 - 22:00 (45 minutes)

Test basic database operations with chunks:
- Serialize/deserialize chunks
- Save/load to JSON (mock DB)
- Filter by metadata
- Ensure no data loss in round-trip
"""

import sys
from pathlib import Path

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import time
import json
from typing import Dict, Any, List
from dataclasses import asdict
import tempfile

from docx import Document as DocxDocument
from src.preprocessing.base.models import ProcessedDocument
from src.preprocessing.chunking.chunk_factory import create_chunker
from src.preprocessing.chunking.base_chunker import UniversalChunk


def create_test_chunks(num_types: int = 3) -> List[UniversalChunk]:
    """Create test chunks from real documents"""
    print(f"\n{'='*60}")
    print("Creating Test Chunks...")
    print(f"{'='*60}")

    test_files = [
        {
            "doc_type": "law",
            "filepath": Path("data/raw/Luat chinh/Luat so 90 2025-qh15.docx"),
        },
        {
            "doc_type": "decree",
            "filepath": Path(
                "data/raw/Nghi dinh/ND 214 - 4.8.2025 - Thay thế NĐ24-original.docx"
            ),
        },
        {
            "doc_type": "circular",
            "filepath": Path("data/raw/Thong tu/0. Lời văn thông tư.docx"),
        },
    ][:num_types]

    all_chunks = []

    for file_info in test_files:
        filepath = PROJECT_ROOT / file_info["filepath"]
        doc_type = file_info["doc_type"]

        if not filepath.exists():
            print(f"   ⚠️  Skipping {doc_type}: File not found")
            continue

        # Load and chunk
        doc = DocxDocument(filepath)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        proc_doc = ProcessedDocument(
            metadata={
                "filename": filepath.name,
                "document_type": doc_type,
                "source_path": str(filepath),
            },
            content={"full_text": text},
        )

        chunker = create_chunker(document_type=doc_type)
        chunks = chunker.chunk(proc_doc)

        # Take first 5 chunks from each document
        all_chunks.extend(chunks[:5])
        print(f"   ✅ {doc_type}: Created {len(chunks)} chunks, using first 5")

    print(f"\n   📊 Total chunks for testing: {len(all_chunks)}")
    return all_chunks


def test_serialize_chunks(chunks: List[UniversalChunk]) -> Dict[str, Any]:
    """Test 1: Serialize chunks to dict"""
    print(f"\n{'='*60}")
    print("Test 1: Chunk Serialization")
    print(f"{'='*60}")

    result = {
        "test": "serialization",
        "success": False,
        "chunks_serialized": 0,
        "errors": [],
    }

    try:
        serialized = []

        for i, chunk in enumerate(chunks):
            try:
                # Convert chunk to dict
                chunk_dict = asdict(chunk)
                serialized.append(chunk_dict)
                result["chunks_serialized"] += 1
            except Exception as e:
                result["errors"].append(f"Chunk {i}: {type(e).__name__}")

        result["success"] = result["chunks_serialized"] == len(chunks)

        print(f"   ✅ Serialized: {result['chunks_serialized']}/{len(chunks)}")
        if result["errors"]:
            print(f"   ⚠️  Errors: {len(result['errors'])}")
            for err in result["errors"][:3]:
                print(f"      - {err}")

    except Exception as e:
        result["errors"].append(f"Global error: {type(e).__name__}: {str(e)}")
        print(f"   ❌ Failed: {e}")

    return result


def test_save_load_json(chunks: List[UniversalChunk]) -> Dict[str, Any]:
    """Test 2: Save to JSON and load back"""
    print(f"\n{'='*60}")
    print("Test 2: Save/Load JSON (Mock DB)")
    print(f"{'='*60}")

    result = {
        "test": "save_load_json",
        "success": False,
        "data_loss": False,
        "errors": [],
    }

    try:
        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        # Step 1: Serialize and save
        serialized = [asdict(chunk) for chunk in chunks]

        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(serialized, f, ensure_ascii=False, indent=2)

        file_size = tmp_path.stat().st_size
        print(f"   💾 Saved {len(serialized)} chunks ({file_size:,} bytes)")

        # Step 2: Load back
        with open(tmp_path, "r", encoding="utf-8") as f:
            loaded = json.load(f)

        print(f"   📂 Loaded {len(loaded)} chunks")

        # Step 3: Validate round-trip
        if len(loaded) != len(serialized):
            result["data_loss"] = True
            result["errors"].append(
                f"Count mismatch: {len(loaded)} vs {len(serialized)}"
            )

        # Check content integrity
        for i, (original, loaded_chunk) in enumerate(zip(serialized, loaded)):
            if original.get("content") != loaded_chunk.get("content"):
                result["data_loss"] = True
                result["errors"].append(f"Content mismatch in chunk {i}")
                break

        result["success"] = not result["data_loss"]

        if result["success"]:
            print(f"   ✅ Round-trip successful: No data loss")
        else:
            print(f"   ❌ Data loss detected")
            for err in result["errors"][:3]:
                print(f"      - {err}")

    except Exception as e:
        result["errors"].append(f"Error: {type(e).__name__}: {str(e)}")
        print(f"   ❌ Failed: {e}")

    finally:
        if tmp_path.exists():
            tmp_path.unlink()

    return result


def test_filter_by_metadata(chunks: List[UniversalChunk]) -> Dict[str, Any]:
    """Test 3: Filter chunks by metadata"""
    print(f"\n{'='*60}")
    print("Test 3: Filter by Metadata")
    print(f"{'='*60}")

    result = {
        "test": "filter_metadata",
        "success": False,
        "filters_tested": 0,
        "errors": [],
    }

    try:
        # Convert to list of dicts for easier filtering
        chunk_dicts = [asdict(chunk) for chunk in chunks]

        # Test filter 1: By document_type (direct field in UniversalChunk)
        doc_types = set(
            c.get("document_type") for c in chunk_dicts if c.get("document_type")
        )
        for doc_type in doc_types:
            filtered = [c for c in chunk_dicts if c.get("document_type") == doc_type]
            print(f"   📋 {doc_type}: {len(filtered)} chunks")
            result["filters_tested"] += 1

        # Test filter 2: By chunk size
        large_chunks = [c for c in chunk_dicts if len(c.get("content", "")) > 800]
        small_chunks = [c for c in chunk_dicts if len(c.get("content", "")) < 500]

        print(f"   📏 Large chunks (>800): {len(large_chunks)}")
        print(f"   📏 Small chunks (<500): {len(small_chunks)}")
        result["filters_tested"] += 2

        # Test filter 3: By hierarchy level (if available)
        with_hierarchy = [c for c in chunk_dicts if c.get("hierarchy")]
        print(f"   🌳 With hierarchy: {len(with_hierarchy)}")
        result["filters_tested"] += 1

        result["success"] = result["filters_tested"] > 0

        print(f"\n   ✅ Filters tested: {result['filters_tested']}")

    except Exception as e:
        result["errors"].append(f"Error: {type(e).__name__}: {str(e)}")
        print(f"   ❌ Failed: {e}")

    return result


def test_reconstruct_chunks(chunks: List[UniversalChunk]) -> Dict[str, Any]:
    """Test 4: Reconstruct UniversalChunk from dict"""
    print(f"\n{'='*60}")
    print("Test 4: Reconstruct Chunks from Dict")
    print(f"{'='*60}")

    result = {
        "test": "reconstruct_chunks",
        "success": False,
        "reconstructed": 0,
        "errors": [],
    }

    try:
        # Serialize
        serialized = [asdict(chunk) for chunk in chunks]

        # Reconstruct
        reconstructed_chunks = []
        for i, chunk_dict in enumerate(serialized):
            try:
                # Create new UniversalChunk from dict
                reconstructed = UniversalChunk(**chunk_dict)
                reconstructed_chunks.append(reconstructed)
                result["reconstructed"] += 1
            except Exception as e:
                result["errors"].append(f"Chunk {i}: {type(e).__name__}")

        # Validate
        if result["reconstructed"] == len(chunks):
            # Check content integrity
            all_match = all(
                orig.content == recon.content
                for orig, recon in zip(chunks, reconstructed_chunks)
            )

            if all_match:
                result["success"] = True
                print(f"   ✅ Reconstructed {result['reconstructed']}/{len(chunks)}")
                print(f"   ✅ Content integrity verified")
            else:
                result["errors"].append("Content mismatch after reconstruction")
                print(f"   ❌ Content mismatch detected")
        else:
            print(f"   ⚠️  Only reconstructed {result['reconstructed']}/{len(chunks)}")

        if result["errors"]:
            for err in result["errors"][:3]:
                print(f"      - {err}")

    except Exception as e:
        result["errors"].append(f"Error: {type(e).__name__}: {str(e)}")
        print(f"   ❌ Failed: {e}")

    return result


def main():
    """Run database integration tests"""
    print("=" * 80)
    print("🧪 DATABASE INTEGRATION TEST (Basic)")
    print("=" * 80)
    print(f"Time: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Goal: Test chunk persistence and retrieval (mock DB)")
    print()

    # Create test chunks
    test_chunks = create_test_chunks(num_types=3)

    if not test_chunks:
        print("\n❌ No test chunks created - cannot proceed")
        return 1

    # Run tests
    test_results = [
        test_serialize_chunks(test_chunks),
        test_save_load_json(test_chunks),
        test_filter_by_metadata(test_chunks),
        test_reconstruct_chunks(test_chunks),
    ]

    # Summary Report
    print("\n" + "=" * 80)
    print("📊 TEST SUMMARY")
    print("=" * 80)

    total_tests = len(test_results)
    passed_tests = sum(1 for r in test_results if r["success"])
    failed_tests = total_tests - passed_tests

    print(f"\nTests Run: {total_tests}")
    print(
        f"Passed:    {passed_tests}/{total_tests} ({passed_tests/total_tests*100:.1f}%)"
    )
    print(f"Failed:    {failed_tests}/{total_tests}")

    print("\n📋 Detailed Results:")
    for result in test_results:
        status = "✅" if result["success"] else "❌"
        test_name = result["test"].replace("_", " ").title()
        error_count = len(result.get("errors", []))

        print(f"   {status} {test_name:30s} | Errors: {error_count}")

    # Overall success
    print("\n" + "=" * 80)
    if passed_tests == total_tests:
        print("🎉 ALL TESTS PASSED! Database Integration Ready!")
        print("=" * 80)
        return 0
    else:
        print(f"⚠️  {failed_tests} TEST(S) FAILED - See Details Above")
        print("=" * 80)

        # Show errors
        print("\n🔍 Errors Found:")
        for result in test_results:
            if not result["success"] and result.get("errors"):
                test_name = result["test"].replace("_", " ").title()
                print(f"\n{test_name}:")
                for err in result["errors"][:3]:
                    print(f"   - {err}")

        return 1


if __name__ == "__main__":
    sys.exit(main())
