#!/usr/bin/env python3
"""
Run full chunking quality test and generate updated report.
This is a standalone version that doesn't use pytest.
"""

import sys
import time
import json
from pathlib import Path
from collections import defaultdict
from datetime import datetime

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document as DocxDocument
from src.preprocessing.chunking.bidding_hybrid_chunker import BiddingHybridChunker
from src.preprocessing.chunking.hierarchical_chunker import HierarchicalChunker
from src.preprocessing.chunking.report_hybrid_chunker import ReportHybridChunker
from src.preprocessing.base.models import ProcessedDocument


def load_docx_to_processed_doc(filepath: Path, doc_type: str) -> ProcessedDocument:
    """Load DOCX file and convert to ProcessedDocument."""
    doc = DocxDocument(filepath)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    return ProcessedDocument(
        metadata={
            "filename": filepath.name,
            "document_type": doc_type,
        },
        content={
            "full_text": text,
        },
    )


def select_chunker(doc_type: str):
    """Select appropriate chunker for document type."""
    if doc_type == "bidding":
        return BiddingHybridChunker(
            min_size=300,
            max_size=1500,
            target_size=800,
        )
    elif doc_type == "report":
        return ReportHybridChunker(
            min_size=400,  # More aggressive merge
            max_size=1500,
            target_size=800,
        )
    else:  # law, decree, circular, decision
        return HierarchicalChunker(
            min_size=300,
            max_size=1500,
            target_size=800,
        )


def test_document_type(doc_type: str, file_paths: list) -> dict:
    """Test all documents of a specific type."""
    print(f"\n📊 {doc_type.upper()}")
    print("─" * 80)

    chunker = select_chunker(doc_type)

    results = {
        "doc_type": doc_type,
        "total_files": len(file_paths),
        "loaded_files": 0,
        "failed_files": [],
        "total_chunks": 0,
        "in_range_chunks": 0,
        "chunk_sizes": [],
        "file_details": [],
    }

    for filepath in sorted(file_paths):
        try:
            # Load document
            proc_doc = load_docx_to_processed_doc(filepath, doc_type)
            results["loaded_files"] += 1

            # Chunk document
            chunks = chunker.chunk(proc_doc)

            # Calculate quality
            chunks_in_range = sum(1 for c in chunks if 300 <= len(c.content) <= 1500)
            quality = chunks_in_range / len(chunks) * 100 if chunks else 0

            results["total_chunks"] += len(chunks)
            results["in_range_chunks"] += chunks_in_range
            results["chunk_sizes"].extend([len(c.content) for c in chunks])

            status = "✅" if quality >= 80 else "⚠️"
            print(
                f"   {status} {filepath.name:50s} | {len(chunks):3d} chunks | {chunks_in_range:3d}/{len(chunks):3d} in-range ({quality:5.1f}%) | avg {sum(len(c.content) for c in chunks) // len(chunks):4d} chars"
            )

            results["file_details"].append(
                {
                    "filename": filepath.name,
                    "chunks": len(chunks),
                    "in_range": chunks_in_range,
                    "quality": quality,
                    "avg_size": (
                        sum(len(c.content) for c in chunks) // len(chunks)
                        if chunks
                        else 0
                    ),
                }
            )

        except Exception as e:
            results["failed_files"].append(
                {
                    "filename": filepath.name,
                    "error": f"{type(e).__name__}: {str(e)[:100]}",
                }
            )
            print(f"   ❌ {filepath.name:50s} | ERROR: {type(e).__name__}")

    # Calculate summary stats
    quality = (
        results["in_range_chunks"] / results["total_chunks"] * 100
        if results["total_chunks"] > 0
        else 0
    )

    print(
        f"\n   📋 Summary: {results['loaded_files']}/{results['total_files']} files loaded"
    )
    print(
        f"   📊 Quality: {results['in_range_chunks']}/{results['total_chunks']} chunks in-range ({quality:.1f}%)"
    )
    if results["chunk_sizes"]:
        print(
            f"   📏 Size: {min(results['chunk_sizes'])}-{max(results['chunk_sizes'])} chars, avg {sum(results['chunk_sizes'])//len(results['chunk_sizes'])} chars"
        )

    return results


def main():
    """Run full quality test."""
    print("=" * 80)
    print("CHUNKING QUALITY TEST - FULL UPDATE")
    print("=" * 80)
    print(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    start_time = time.time()

    # Define document collections
    doc_collections = {
        "bidding": list(
            Path(
                "data/raw/Ho so moi thau/1. Kế hoạch tổng thể LCNT_kế hoạch LCNT"
            ).glob("*.docx")
        ),
        "law": list(Path("data/raw/Luat chinh").glob("*.docx")),
        "decree": list(Path("data/raw/Nghi dinh").glob("*.docx")),
        "circular": list(Path("data/raw/Thong tu").glob("*.docx")),
        "decision": list(Path("data/raw/Quyet dinh").glob("*.docx")),
        "report": list(
            Path("data/raw/Mau bao cao/14. Báo cáo đánh giá").glob("*.docx")
        ),
    }

    all_results = {}

    for doc_type, file_paths in doc_collections.items():
        if file_paths:
            all_results[doc_type] = test_document_type(doc_type, file_paths)

    # Overall summary
    print("\n" + "=" * 80)
    print("OVERALL SUMMARY")
    print("=" * 80)

    total_files = sum(r["total_files"] for r in all_results.values())
    loaded_files = sum(r["loaded_files"] for r in all_results.values())
    total_chunks = sum(r["total_chunks"] for r in all_results.values())
    in_range_chunks = sum(r["in_range_chunks"] for r in all_results.values())

    overall_quality = in_range_chunks / total_chunks * 100 if total_chunks > 0 else 0

    print(f"\nDocument types tested: {len(all_results)}")
    print(f"Total files: {total_files}")
    print(
        f"Files loaded: {loaded_files}/{total_files} ({loaded_files/total_files*100:.1f}%)"
    )
    print(f"Total chunks: {total_chunks:,}")
    print(
        f"In-range chunks: {in_range_chunks:,}/{total_chunks:,} ({overall_quality:.1f}%)"
    )
    print(f"\nExecution time: {time.time() - start_time:.2f} seconds")

    # Save results to JSON
    output_file = Path("data/outputs/chunking_metrics_updated.json")
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(
            {
                "timestamp": datetime.now().isoformat(),
                "overall": {
                    "total_files": total_files,
                    "loaded_files": loaded_files,
                    "total_chunks": total_chunks,
                    "in_range_chunks": in_range_chunks,
                    "quality": overall_quality,
                },
                "by_type": all_results,
            },
            f,
            indent=2,
            ensure_ascii=False,
        )

    print(f"\n✅ Results saved to: {output_file}")
    print("=" * 80)

    return overall_quality >= 80


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
