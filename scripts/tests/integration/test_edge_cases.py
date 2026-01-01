#!/usr/bin/env python3
"""
Test #3: Edge Cases & Error Handling
Timeline: 20:30 - 21:15 (45 minutes)

Test system behavior with problematic inputs:
- Empty/minimal documents
- Very large documents
- Malformed files
- Special characters
- Missing required fields
"""

import sys
from pathlib import Path

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import time
from typing import Dict, Any, List
import tempfile

from docx import Document as DocxDocument
from docx.shared import Pt
from src.preprocessing.base.models import ProcessedDocument
from src.preprocessing.chunking.chunk_factory import create_chunker


def create_test_docx(content: str, output_path: Path) -> None:
    """Create a test DOCX file with given content"""
    doc = DocxDocument()
    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(output_path)


def test_empty_document() -> Dict[str, Any]:
    """Test with completely empty document"""
    print(f"\n{'='*60}")
    print("Test 1: Empty Document")
    print(f"{'='*60}")

    result = {
        "test": "empty_document",
        "success": False,
        "error_type": None,
        "error_message": None,
        "chunks_created": 0,
    }

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        # Create empty document
        create_test_docx("", tmp_path)

        # Try to process
        proc_doc = ProcessedDocument(
            metadata={"filename": "empty.docx", "document_type": "law"},
            content={"full_text": ""},
        )

        chunker = create_chunker(document_type="law")
        chunks = chunker.chunk(proc_doc)

        result["chunks_created"] = len(chunks)
        result["success"] = True

        print(f"   ✅ Handled gracefully: {len(chunks)} chunks created")

    except Exception as e:
        result["error_type"] = type(e).__name__
        result["error_message"] = str(e)
        print(f"   ⚠️  Error: {type(e).__name__}: {str(e)[:60]}")

    finally:
        if tmp_path.exists():
            tmp_path.unlink()

    return result


def test_minimal_document() -> Dict[str, Any]:
    """Test with minimal valid content"""
    print(f"\n{'='*60}")
    print("Test 2: Minimal Document (< 100 chars)")
    print(f"{'='*60}")

    result = {
        "test": "minimal_document",
        "success": False,
        "error_type": None,
        "error_message": None,
        "chunks_created": 0,
    }

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        # Create minimal document
        content = "Điều 1. Quy định chung\nVăn bản này có hiệu lực."
        create_test_docx(content, tmp_path)

        proc_doc = ProcessedDocument(
            metadata={"filename": "minimal.docx", "document_type": "law"},
            content={"full_text": content},
        )

        chunker = create_chunker(document_type="law")
        chunks = chunker.chunk(proc_doc)

        result["chunks_created"] = len(chunks)
        result["success"] = True

        print(f"   ✅ Created {len(chunks)} chunk(s)")
        if chunks:
            print(f"   📏 Chunk size: {len(chunks[0].content)} chars")

    except Exception as e:
        result["error_type"] = type(e).__name__
        result["error_message"] = str(e)
        print(f"   ⚠️  Error: {type(e).__name__}: {str(e)[:60]}")

    finally:
        if tmp_path.exists():
            tmp_path.unlink()

    return result


def test_large_document() -> Dict[str, Any]:
    """Test with very large document (> 1MB text)"""
    print(f"\n{'='*60}")
    print("Test 3: Large Document (> 1MB)")
    print(f"{'='*60}")

    result = {
        "test": "large_document",
        "success": False,
        "error_type": None,
        "error_message": None,
        "chunks_created": 0,
        "execution_time": 0,
    }

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        # Create large document (>1MB text)
        large_content = "\n".join(
            [
                f"Điều {i}. Quy định về vấn đề số {i}\n"
                f"Nội dung chi tiết của điều luật số {i}. " * 50
                for i in range(1, 1001)  # 1000 articles
            ]
        )

        print(f"   📝 Generating {len(large_content):,} chars...")
        create_test_docx(large_content, tmp_path)

        proc_doc = ProcessedDocument(
            metadata={"filename": "large.docx", "document_type": "law"},
            content={"full_text": large_content},
        )

        start_time = time.time()
        chunker = create_chunker(document_type="law")
        chunks = chunker.chunk(proc_doc)
        result["execution_time"] = time.time() - start_time

        result["chunks_created"] = len(chunks)
        result["success"] = True

        print(f"   ✅ Created {len(chunks)} chunks in {result['execution_time']:.2f}s")
        print(f"   📊 Speed: {len(chunks)/result['execution_time']:.0f} chunks/sec")

    except Exception as e:
        result["error_type"] = type(e).__name__
        result["error_message"] = str(e)
        print(f"   ⚠️  Error: {type(e).__name__}: {str(e)[:60]}")

    finally:
        if tmp_path.exists():
            tmp_path.unlink()

    return result


def test_special_characters() -> Dict[str, Any]:
    """Test with special characters and Unicode"""
    print(f"\n{'='*60}")
    print("Test 4: Special Characters & Unicode")
    print(f"{'='*60}")

    result = {
        "test": "special_characters",
        "success": False,
        "error_type": None,
        "error_message": None,
        "chunks_created": 0,
    }

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        # Create document with special chars
        content = """
Điều 1. Quy định về ký tự đặc biệt
1. Ký tự Việt: àáảãạ ÀÁẢÃẠ êếệễề ÊẾỆỄỀ
2. Ký hiệu: @#$%^&*()[]{}|\\/<>?!~`
3. Unicode: 🎉 💯 ✅ ❌ ⚠️ 📊
4. Số học: ≤ ≥ ≠ ± × ÷
5. Ngoặc kép: "abc" 'xyz' «test» ‹single›
        """.strip()

        create_test_docx(content, tmp_path)

        proc_doc = ProcessedDocument(
            metadata={"filename": "special.docx", "document_type": "law"},
            content={"full_text": content},
        )

        chunker = create_chunker(document_type="law")
        chunks = chunker.chunk(proc_doc)

        result["chunks_created"] = len(chunks)
        result["success"] = True

        print(f"   ✅ Created {len(chunks)} chunk(s)")
        if chunks:
            has_unicode = any(ord(c) > 127 for c in chunks[0].content)
            print(f"   📝 Unicode preserved: {has_unicode}")

    except Exception as e:
        result["error_type"] = type(e).__name__
        result["error_message"] = str(e)
        print(f"   ⚠️  Error: {type(e).__name__}: {str(e)[:60]}")

    finally:
        if tmp_path.exists():
            tmp_path.unlink()

    return result


def test_malformed_metadata() -> Dict[str, Any]:
    """Test with missing/invalid metadata"""
    print(f"\n{'='*60}")
    print("Test 5: Malformed Metadata")
    print(f"{'='*60}")

    result = {
        "test": "malformed_metadata",
        "success": False,
        "error_type": None,
        "error_message": None,
        "chunks_created": 0,
    }

    try:
        # Missing document_type
        proc_doc = ProcessedDocument(
            metadata={"filename": "test.docx"},  # Missing document_type!
            content={"full_text": "Điều 1. Test content here."},
        )

        # Should fail gracefully
        chunker = create_chunker(document_type="law")
        chunks = chunker.chunk(proc_doc)

        result["chunks_created"] = len(chunks)
        result["success"] = True

        print(f"   ✅ Handled missing metadata: {len(chunks)} chunks")

    except Exception as e:
        result["error_type"] = type(e).__name__
        result["error_message"] = str(e)

        # Check if error message is clear
        is_clear = "metadata" in str(e).lower() or "document_type" in str(e).lower()

        if is_clear:
            print(f"   ✅ Clear error message: {str(e)[:60]}")
            result["success"] = True  # Good error handling
        else:
            print(f"   ❌ Unclear error: {str(e)[:60]}")

    return result


def test_invalid_document_type() -> Dict[str, Any]:
    """Test with invalid document type"""
    print(f"\n{'='*60}")
    print("Test 6: Invalid Document Type")
    print(f"{'='*60}")

    result = {
        "test": "invalid_doc_type",
        "success": False,
        "error_type": None,
        "error_message": None,
        "chunks_created": 0,
    }

    try:
        proc_doc = ProcessedDocument(
            metadata={
                "filename": "test.docx",
                "document_type": "invalid_type_xyz",  # Invalid!
            },
            content={"full_text": "Điều 1. Test content here."},
        )

        # Try to create chunker with invalid type
        chunker = create_chunker(document_type="invalid_type_xyz")
        chunks = chunker.chunk(proc_doc)

        result["chunks_created"] = len(chunks)
        result["success"] = True

        print(f"   ✅ Fallback chunker used: {len(chunks)} chunks")
        print(f"   📝 Chunker: {chunker.__class__.__name__}")

    except Exception as e:
        result["error_type"] = type(e).__name__
        result["error_message"] = str(e)

        # Check if error is informative
        is_clear = "type" in str(e).lower() or "invalid" in str(e).lower()

        if is_clear:
            print(f"   ✅ Clear error: {str(e)[:60]}")
            result["success"] = True
        else:
            print(f"   ❌ Unclear error: {str(e)[:60]}")

    return result


def main():
    """Run edge case tests"""
    print("=" * 80)
    print("🧪 EDGE CASES & ERROR HANDLING INTEGRATION TEST")
    print("=" * 80)
    print(f"Time: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Goal: Test system behavior with problematic inputs")
    print()

    # Run all tests
    test_results = [
        test_empty_document(),
        test_minimal_document(),
        test_large_document(),
        test_special_characters(),
        test_malformed_metadata(),
        test_invalid_document_type(),
    ]

    # Summary Report
    print("\n" + "=" * 80)
    print("📊 TEST SUMMARY")
    print("=" * 80)

    total_tests = len(test_results)
    passed_tests = sum(1 for r in test_results if r["success"])
    failed_tests = total_tests - passed_tests

    print(f"\nTests Run: {total_tests}")
    print(
        f"Passed:    {passed_tests}/{total_tests} ({passed_tests/total_tests*100:.1f}%)"
    )
    print(f"Failed:    {failed_tests}/{total_tests}")

    print("\n📋 Detailed Results:")
    for result in test_results:
        status = "✅" if result["success"] else "❌"
        test_name = result["test"].replace("_", " ").title()

        if result["error_type"]:
            print(f"   {status} {test_name:25s} | Error: {result['error_type']}")
        else:
            chunks = result.get("chunks_created", 0)
            print(f"   {status} {test_name:25s} | Chunks: {chunks}")

    # Error Handling Quality
    print("\n🛡️  Error Handling Quality:")
    errors_handled = sum(1 for r in test_results if r["error_type"] and r["success"])
    errors_total = sum(1 for r in test_results if r["error_type"])

    if errors_total > 0:
        print(f"   Errors handled gracefully: {errors_handled}/{errors_total}")
    else:
        print(f"   No errors encountered (robust system!)")

    # Overall success
    print("\n" + "=" * 80)
    if passed_tests == total_tests:
        print("🎉 ALL TESTS PASSED! Excellent Error Handling!")
        print("=" * 80)
        return 0
    else:
        print(f"⚠️  {failed_tests} TEST(S) FAILED - Needs Improvement")
        print("=" * 80)

        # Show failures
        print("\n🔍 Failed Tests:")
        for result in test_results:
            if not result["success"]:
                test_name = result["test"].replace("_", " ").title()
                msg = result.get("error_message", "Unknown error")
                print(f"   - {test_name}: {msg[:60]}")

        return 1


if __name__ == "__main__":
    sys.exit(main())
