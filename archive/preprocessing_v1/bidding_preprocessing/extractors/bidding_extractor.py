"""
Bidding Document Extractor

Extracts content from Vietnamese bidding documents (Hồ sơ mời thầu).
Supports various document formats including DOC, DOCX for different bidding types.
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")

try:
    import win32com.client

    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


@dataclass
class ExtractedContent:
    """Container for extracted bidding document content"""

    text: str
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]]
    statistics: Dict[str, Any]


class BiddingExtractor:
    """
    Extractor for Vietnamese bidding documents (Hồ sơ mời thầu)
    Handles various bidding document types and formats
    """

    def __init__(self, preserve_formatting: bool = True):
        """
        Initialize bidding document extractor

        Args:
            preserve_formatting: Whether to preserve document formatting
        """
        self.preserve_formatting = preserve_formatting

        # Bidding document type patterns
        self.bidding_patterns = {
            "construction": r"(?i)(xây\s*lắp|xây\s*dựng|thi\s*công)",
            "goods": r"(?i)(hàng\s*hóa|thiết\s*bị|vật\s*tư)",
            "consulting": r"(?i)(tư\s*vấn|dịch\s*vụ\s*tư\s*vấn)",
            "non_consulting": r"(?i)(phi\s*tư\s*vấn|dịch\s*vụ\s*phi\s*tư\s*vấn)",
            "epc": r"(?i)(epc|engineering.*procurement.*construction)",
            "ep": r"(?i)(ep|engineering.*procurement)",
            "ec": r"(?i)(ec|engineering.*construction)",
            "pc": r"(?i)(pc|procurement.*construction)",
            "equipment_lease": r"(?i)(máy\s*đặt|máy\s*mượn|thuê\s*thiết\s*bị)",
            "online_bidding": r"(?i)(chào\s*giá\s*trực\s*tuyến|đấu\s*thầu\s*trực\s*tuyến)",
            "online_procurement": r"(?i)(mua\s*sắm\s*trực\s*tuyến)",
        }

        # Bidding document structure patterns
        self.structure_patterns = {
            "section": r"^(PHẦN|Phần)\s+([IVXLCDM]+|\d+)[\.\:]?\s*(.+)$",
            "chapter": r"^(CHƯƠNG|Chương)\s+([IVXLCDM]+|\d+)[\.\:]?\s*(.+)$",
            "article": r"^(Điều|ĐIỀU)\s+(\d+[a-z]?)[\.\:]?\s*(.+)$",
            "clause": r"^(\d+)[\.\)]?\s+(.+)",
            "point": r"^([a-zđ])\)\s+(.+)",
            "sub_point": r"^-\s+(.+)",
            "table_title": r"^(Bảng|BẢNG)\s+(\d+)[\.\:]?\s*(.+)$",
            "form_title": r"^(Mẫu|MẪU)\s+(\d+[A-Z]?)[\.\:]?\s*(.+)$",
            "appendix": r"^(Phụ\s*lục|PHỤ\s*LỤC)\s+(\d+[A-Z]?)[\.\:]?\s*(.+)$",
        }

        # Bidding terminology patterns
        self.bidding_terms = {
            "invitation": ["mời thầu", "mời chào giá", "mời quan tâm"],
            "contractor": ["nhà thầu", "đơn vị thầu", "thầu phụ"],
            "owner": ["chủ đầu tư", "bên mời thầu", "chủ dự án"],
            "contract": ["hợp đồng", "gói thầu", "dự án"],
            "evaluation": ["đánh giá", "xét thầu", "chấm thầu"],
            "technical": ["kỹ thuật", "chuyên môn", "công nghệ"],
            "financial": ["tài chính", "kinh tế", "chi phí"],
            "qualification": ["năng lực", "kinh nghiệm", "trình độ"],
        }

    def extract(self, file_path: Path) -> ExtractedContent:
        """
        Extract content from bidding document file

        Args:
            file_path: Path to bidding document file

        Returns:
            ExtractedContent with text, metadata, tables, and statistics
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Determine extraction method based on file extension
        if file_path.suffix.lower() == ".docx":
            return self._extract_docx(file_path)
        elif file_path.suffix.lower() == ".doc":
            if WORD_AVAILABLE:
                return self._extract_doc_with_word(file_path)
            else:
                raise RuntimeError(
                    "python-docx cannot handle .doc files. Install pywin32 for Word COM support"
                )
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

    def _extract_docx(self, file_path: Path) -> ExtractedContent:
        """Extract content from DOCX file"""
        try:
            doc = Document(file_path)

            # Extract text content
            text_parts = []
            tables = []
            structure_elements = 0

            for element in doc.element.body:
                if element.tag.endswith("p"):  # Paragraph
                    para = Paragraph(element, doc)
                    if para.text.strip():
                        # Check if paragraph matches structure patterns
                        if self._is_structure_element(para.text):
                            structure_elements += 1
                        text_parts.append(para.text.strip())

                elif element.tag.endswith("tbl"):  # Table
                    table = Table(element, doc)
                    table_data = self._extract_table_data(table)
                    if table_data:
                        tables.append(table_data)
                        # Add table content to text
                        text_parts.append(
                            f"[BẢNG: {table_data.get('title', 'Không tiêu đề')}]"
                        )

            full_text = "\n\n".join(text_parts)

            # Extract metadata
            metadata = self._extract_metadata(file_path, full_text)

            # Calculate statistics
            statistics = {
                "original_chars": len(full_text),
                "paragraphs": len([p for p in text_parts if p.strip()]),
                "tables": len(tables),
                "structure_elements": structure_elements,
                "bidding_type": self._detect_bidding_type(full_text),
                "contains_forms": self._contains_forms(full_text),
                "contains_technical_specs": self._contains_technical_specs(full_text),
            }

            return ExtractedContent(
                text=full_text, metadata=metadata, tables=tables, statistics=statistics
            )

        except Exception as e:
            raise RuntimeError(f"Error extracting DOCX file {file_path}: {str(e)}")

    def _extract_doc_with_word(self, file_path: Path) -> ExtractedContent:
        """Extract content from DOC file using Word COM"""
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(file_path))

            # Extract text
            full_text = doc.Content.Text

            # Extract tables
            tables = []
            for i, table in enumerate(doc.Tables):
                table_data = {
                    "id": i,
                    "title": f"Bảng {i+1}",
                    "rows": table.Rows.Count,
                    "columns": table.Columns.Count,
                    "content": [],
                }

                try:
                    for row in range(1, min(table.Rows.Count + 1, 6)):  # First 5 rows
                        row_data = []
                        for col in range(1, table.Columns.Count + 1):
                            try:
                                cell_text = table.Cell(row, col).Range.Text.strip()
                                row_data.append(cell_text.replace("\r\x07", ""))
                            except:
                                row_data.append("")
                        table_data["content"].append(row_data)
                except:
                    pass

                tables.append(table_data)

            doc.Close()
            word.Quit()

            # Extract metadata and statistics
            metadata = self._extract_metadata(file_path, full_text)
            statistics = {
                "original_chars": len(full_text),
                "tables": len(tables),
                "bidding_type": self._detect_bidding_type(full_text),
                "contains_forms": self._contains_forms(full_text),
                "contains_technical_specs": self._contains_technical_specs(full_text),
            }

            return ExtractedContent(
                text=full_text, metadata=metadata, tables=tables, statistics=statistics
            )

        except Exception as e:
            raise RuntimeError(f"Error extracting DOC file {file_path}: {str(e)}")

    def _extract_table_data(self, table: Table) -> Dict[str, Any]:
        """Extract data from a table"""
        table_data = {
            "id": id(table),
            "title": "Không tiêu đề",
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0,
            "content": [],
        }

        try:
            for i, row in enumerate(table.rows[:5]):  # First 5 rows
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace("\n", " ")
                    row_data.append(cell_text)
                table_data["content"].append(row_data)

                # Try to extract table title from first row
                if i == 0 and len(row_data) > 0:
                    first_cell = row_data[0]
                    if any(
                        keyword in first_cell.lower()
                        for keyword in ["bảng", "mẫu", "danh sách"]
                    ):
                        table_data["title"] = first_cell[:100]
        except:
            pass

        return table_data

    def _is_structure_element(self, text: str) -> bool:
        """Check if text is a structure element"""
        for pattern in self.structure_patterns.values():
            if re.match(pattern, text.strip(), re.MULTILINE):
                return True
        return False

    def _extract_metadata(self, file_path: Path, text: str) -> Dict[str, Any]:
        """Extract metadata from bidding document"""
        metadata = {
            "source_file": file_path.name,
            "file_type": "bidding_document",
            "document_category": self._detect_bidding_type(text),
            "extracted_at": datetime.now().isoformat(),
            "file_size_bytes": file_path.stat().st_size,
        }

        # Extract bidding-specific metadata
        metadata.update(self._extract_bidding_metadata(text))

        return metadata

    def _extract_bidding_metadata(self, text: str) -> Dict[str, Any]:
        """Extract bidding-specific metadata"""
        metadata = {}

        # Extract project information
        project_patterns = [
            r"(?i)dự\s*án[:\s]*(.+?)(?:\n|$)",
            r"(?i)tên\s*dự\s*án[:\s]*(.+?)(?:\n|$)",
            r"(?i)project[:\s]*(.+?)(?:\n|$)",
        ]

        for pattern in project_patterns:
            match = re.search(pattern, text)
            if match:
                metadata["project_name"] = match.group(1).strip()[:200]
                break

        # Extract contract package information
        package_patterns = [
            r"(?i)gói\s*thầu[:\s]*(.+?)(?:\n|$)",
            r"(?i)package[:\s]*(.+?)(?:\n|$)",
            r"(?i)contract\s*package[:\s]*(.+?)(?:\n|$)",
        ]

        for pattern in package_patterns:
            match = re.search(pattern, text)
            if match:
                metadata["package_name"] = match.group(1).strip()[:200]
                break

        # Extract bidding method
        method_patterns = [
            r"(?i)(đấu\s*thầu\s*rộng\s*rãi)",
            r"(?i)(đấu\s*thầu\s*hạn\s*chế)",
            r"(?i)(chào\s*hàng\s*cạnh\s*tranh)",
            r"(?i)(chỉ\s*định\s*thầu)",
        ]

        for pattern in method_patterns:
            match = re.search(pattern, text)
            if match:
                metadata["bidding_method"] = match.group(1)
                break

        # Extract owner information
        owner_patterns = [
            r"(?i)chủ\s*đầu\s*tư[:\s]*(.+?)(?:\n|$)",
            r"(?i)investor[:\s]*(.+?)(?:\n|$)",
            r"(?i)owner[:\s]*(.+?)(?:\n|$)",
        ]

        for pattern in owner_patterns:
            match = re.search(pattern, text)
            if match:
                metadata["owner"] = match.group(1).strip()[:200]
                break

        return metadata

    def _detect_bidding_type(self, text: str) -> str:
        """Detect the type of bidding document"""
        text_lower = text.lower()

        # Count matches for each type
        type_scores = {}
        for bidding_type, pattern in self.bidding_patterns.items():
            matches = len(re.findall(pattern, text_lower))
            if matches > 0:
                type_scores[bidding_type] = matches

        # Return type with highest score
        if type_scores:
            return max(type_scores.items(), key=lambda x: x[1])[0]

        return "general_bidding"

    def _contains_forms(self, text: str) -> bool:
        """Check if document contains forms/templates"""
        form_indicators = [
            r"(?i)mẫu\s*số",
            r"(?i)biểu\s*mẫu",
            r"(?i)phụ\s*lục",
            r"(?i)form\s*\d+",
            r"\[.*\]",  # Placeholder brackets
            r"\.\.\.\.\.+",  # Dotted lines for filling
        ]

        return any(re.search(pattern, text) for pattern in form_indicators)

    def _contains_technical_specs(self, text: str) -> bool:
        """Check if document contains technical specifications"""
        tech_indicators = [
            r"(?i)thông\s*số\s*kỹ\s*thuật",
            r"(?i)yêu\s*cầu\s*kỹ\s*thuật",
            r"(?i)specification",
            r"(?i)technical\s*requirement",
            r"(?i)tiêu\s*chuẩn\s*kỹ\s*thuật",
        ]

        return any(re.search(pattern, text) for pattern in tech_indicators)

    def export_to_md(self, extracted: ExtractedContent, output_path: Path):
        """Export extracted content to Markdown format"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        md_content = []

        # Header
        md_content.append(
            f"# {extracted.metadata.get('source_file', 'Bidding Document')}"
        )
        md_content.append("")

        # Metadata
        md_content.append("## Document Information")
        md_content.append("")
        for key, value in extracted.metadata.items():
            if key != "source_file":
                md_content.append(f"- **{key.replace('_', ' ').title()}:** {value}")
        md_content.append("")

        # Statistics
        if extracted.statistics:
            md_content.append("## Statistics")
            md_content.append("")
            for key, value in extracted.statistics.items():
                md_content.append(f"- **{key.replace('_', ' ').title()}:** {value}")
            md_content.append("")

        # Tables summary
        if extracted.tables:
            md_content.append("## Tables")
            md_content.append("")
            for i, table in enumerate(extracted.tables):
                md_content.append(f"### Table {i+1}: {table.get('title', 'No title')}")
                md_content.append(f"- Rows: {table.get('rows', 0)}")
                md_content.append(f"- Columns: {table.get('columns', 0)}")
                md_content.append("")

        # Content
        md_content.append("## Content")
        md_content.append("")
        md_content.append(extracted.text)

        # Write to file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_content))
