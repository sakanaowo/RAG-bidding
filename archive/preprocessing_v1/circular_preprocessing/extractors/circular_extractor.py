"""
Circular DOCX Extractor

Extract content from Vietnamese Circular documents (.docx files).
Circulars (Thông tư) are administrative documents with specific structure patterns.
"""

import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


@dataclass
class ExtractedContent:
    """Content extracted from Circular DOCX"""

    text: str
    metadata: Dict
    structure: List[Dict]  # Hierarchical structure
    tables: List[Dict]  # Extracted tables
    statistics: Dict


class CircularExtractor:
    """Extract content from Circular DOCX files"""

    def __init__(self, preserve_formatting: bool = False):
        """
        Args:
            preserve_formatting: Whether to preserve formatting (bold, italic)
        """
        self.preserve_formatting = preserve_formatting

        # Patterns for Vietnamese Circular structure detection
        self.patterns = {
            "phan": r"^(PHẦN THỨ|Phần\s+[IVXLCDM]+|PHẦN\s+[IVXLCDM]+)[:\s.]",
            "chuong": r"^(CHƯƠNG|Chương)\s+([IVXLCDM]+|[0-9]+)[:\s.]",
            "muc": r"^(MỤC|Mục)\s+(\d+)[:\s.]",
            "dieu": r"^Điều\s+(\d+[a-z]?)[:\s.]",
            "khoan": r"^\d+\.\s+",
            "diem": r"^[a-zđ]\)\s+",
            # Circular-specific patterns
            "thong_tu": r"THÔNG TƯ|Thông tư",
            "quy_dinh": r"^Quy định\s+(\d+)[:\s.]",
            "huong_dan": r"^Hướng dẫn\s+(\d+)[:\s.]",
        }

    def extract(self, docx_path: str | Path) -> ExtractedContent:
        """
        Extract content from Circular DOCX file

        Args:
            docx_path: Path to .docx file

        Returns:
            ExtractedContent with text, metadata, structure
        """
        docx_path = Path(docx_path)

        if not docx_path.exists():
            raise FileNotFoundError(f"File not found: {docx_path}")

        if not docx_path.suffix.lower() == ".docx":
            raise ValueError(f"Not a DOCX file: {docx_path}")

        try:
            document = Document(docx_path)
        except Exception as e:
            raise ValueError(f"Cannot open DOCX file: {e}")

        # Extract text with structure preservation
        text_parts = []
        structure = []
        tables = []

        for element in document.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, document)
                text, struct_info = self._extract_paragraph(paragraph)
                if text.strip():
                    text_parts.append(text)
                    if struct_info:
                        structure.append(struct_info)

            elif isinstance(element, CT_Tbl):
                table = Table(element, document)
                table_data = self._extract_table(table)
                if table_data:
                    tables.append(table_data)
                    text_parts.append(
                        f"[TABLE: {table_data.get('summary', 'Table content')}]"
                    )

        full_text = "\n".join(text_parts)

        # Extract metadata
        metadata = self._extract_metadata(document, docx_path, full_text)

        # Calculate statistics
        statistics = self._calculate_statistics(full_text, structure, tables)

        return ExtractedContent(
            text=full_text,
            metadata=metadata,
            structure=structure,
            tables=tables,
            statistics=statistics,
        )

    def _extract_paragraph(self, paragraph: Paragraph) -> Tuple[str, Optional[Dict]]:
        """Extract text and structure info from paragraph"""
        text = paragraph.text.strip()
        if not text:
            return "", None

        # Check for structure patterns
        struct_info = None
        for struct_type, pattern in self.patterns.items():
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                struct_info = {
                    "type": struct_type,
                    "text": text,
                    "match": match.groups() if match.groups() else (text,),
                    "level": self._get_structure_level(struct_type),
                }
                break

        return text, struct_info

    def _extract_table(self, table: Table) -> Optional[Dict]:
        """Extract table data"""
        if not table.rows:
            return None

        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(row_data):  # Skip empty rows
                rows_data.append(row_data)

        if not rows_data:
            return None

        return {
            "rows": rows_data,
            "row_count": len(rows_data),
            "col_count": len(rows_data[0]) if rows_data else 0,
            "summary": f"{len(rows_data)} rows × {len(rows_data[0]) if rows_data else 0} cols",
        }

    def _extract_metadata(self, document: Document, file_path: Path, text: str) -> Dict:
        """Extract metadata from document"""
        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "source": "thuvienphapluat.vn",  # Default source
        }

        # Extract document properties
        props = document.core_properties
        if props.title:
            metadata["title"] = props.title
        else:
            # Extract title from filename or content
            metadata["title"] = self._extract_title_from_content(text) or file_path.stem

        if props.author:
            metadata["author"] = props.author
        if props.created:
            metadata["created_date"] = props.created.isoformat()
        if props.modified:
            metadata["modified_date"] = props.modified.isoformat()

        # Extract circular-specific metadata
        metadata.update(self._extract_circular_metadata(text))

        return metadata

    def _extract_title_from_content(self, text: str) -> Optional[str]:
        """Extract title from document content"""
        lines = text.split("\n")
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if len(line) > 10 and any(
                keyword in line.upper()
                for keyword in ["THÔNG TƯ", "QUY ĐỊNH", "HƯỚNG DẪN"]
            ):
                return line
        return None

    def _extract_circular_metadata(self, text: str) -> Dict:
        """Extract Circular-specific metadata"""
        metadata = {}

        # Extract circular number and year
        patterns = [
            r"Thông tư số\s*(\d+/\d+/[A-Z-]+)\s*ngày\s*(\d{1,2}/\d{1,2}/\d{4})",
            r"THÔNG TƯ\s*(\d+/\d+/[A-Z-]+)",
            r"số\s*(\d+/\d+/[A-Z-]+)",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata["circular_number"] = match.group(1)
                if len(match.groups()) > 1:
                    metadata["issue_date"] = match.group(2)
                break

        # Extract issuing agency
        agency_patterns = [
            r"BỘ\s+([A-ZÀ-Ỹ\s]+)",
            r"TỔNG CỤC\s+([A-ZÀ-Ỹ\s]+)",
            r"CỤC\s+([A-ZÀ-Ỹ\s]+)",
        ]

        for pattern in agency_patterns:
            match = re.search(pattern, text)
            if match:
                metadata["issuing_agency"] = match.group(1).strip()
                break

        return metadata

    def _get_structure_level(self, struct_type: str) -> int:
        """Get hierarchy level for structure type"""
        levels = {
            "phan": 1,
            "chuong": 2,
            "muc": 3,
            "dieu": 4,
            "quy_dinh": 4,  # Same level as dieu
            "huong_dan": 4,
            "khoan": 5,
            "diem": 6,
        }
        return levels.get(struct_type, 999)

    def _calculate_statistics(
        self, text: str, structure: List[Dict], tables: List[Dict]
    ) -> Dict:
        """Calculate extraction statistics"""
        return {
            "char_count": len(text),
            "word_count": len(text.split()),
            "line_count": len(text.split("\n")),
            "paragraph_count": len([p for p in text.split("\n\n") if p.strip()]),
            "structure_elements": len(structure),
            "table_count": len(tables),
            "structure_breakdown": self._count_structure_types(structure),
        }

    def _count_structure_types(self, structure: List[Dict]) -> Dict[str, int]:
        """Count structure elements by type"""
        counts = {}
        for item in structure:
            struct_type = item.get("type", "unknown")
            counts[struct_type] = counts.get(struct_type, 0) + 1
        return counts

    def export_to_md(self, extracted: ExtractedContent, output_path: str | Path):
        """Export extracted content to Markdown format"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        md_lines = []

        # Header
        title = extracted.metadata.get("title", "Circular Document")
        md_lines.append(f"# {title}")
        md_lines.append("")

        # Metadata
        md_lines.append("## Metadata")
        for key, value in extracted.metadata.items():
            if key != "title":
                md_lines.append(f"- **{key}**: {value}")
        md_lines.append("")

        # Statistics
        md_lines.append("## Statistics")
        for key, value in extracted.statistics.items():
            md_lines.append(f"- **{key}**: {value}")
        md_lines.append("")

        # Content
        md_lines.append("## Content")
        md_lines.append(extracted.text)

        # Tables (if any)
        if extracted.tables:
            md_lines.append("\n## Tables")
            for i, table in enumerate(extracted.tables):
                md_lines.append(f"\n### Table {i+1}")
                rows = table.get("rows", [])
                if rows:
                    # Create markdown table
                    md_lines.append("| " + " | ".join(rows[0]) + " |")
                    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                    for row in rows[1:]:
                        md_lines.append("| " + " | ".join(row) + " |")

        output_path.write_text("\n".join(md_lines), encoding="utf-8")
