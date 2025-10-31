"""
DOCX Extractor cho văn bản pháp luật đấu thầu

Module này extract text, tables, và structure từ .docx files
với support cho văn bản pháp luật Việt Nam (Luật, Nghị định, Thông tư)
"""

import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


@dataclass
class ExtractedContent:
    """Nội dung được extract từ DOCX"""

    text: str
    metadata: Dict
    structure: List[Dict]  # Hierarchical structure
    tables: List[Dict]  # Extracted tables
    statistics: Dict


class DocxExtractor:
    """Extract content từ DOCX files cho văn bản pháp luật"""

    def __init__(self, preserve_formatting: bool = False):
        """
        Args:
            preserve_formatting: Có giữ formatting (bold, italic) không
        """
        self.preserve_formatting = preserve_formatting

        # Patterns để detect cấu trúc văn bản pháp luật
        self.patterns = {
            "phan": r"^(PHẦN THỨ|Phần\s+[IVXLCDM]+|PHẦN\s+[IVXLCDM]+)[:\s.]",
            "chuong": r"^(CHƯƠNG|Chương)\s+([IVXLCDM]+|[0-9]+)[:\s.]",
            "muc": r"^(MỤC|Mục)\s+(\d+)[:\s.]",
            "dieu": r"^Điều\s+(\d+[a-z]?)[:\s.]",
            "khoan": r"^\d+\.\s+",
            "diem": r"^[a-zđ]\)\s+",
        }

    def extract(self, docx_path: str | Path) -> ExtractedContent:
        """
        Main method để extract content từ DOCX

        Args:
            docx_path: Path đến file .docx

        Returns:
            ExtractedContent với text, metadata, structure
        """
        docx_path = Path(docx_path)

        if not docx_path.exists():
            raise FileNotFoundError(f"File not found: {docx_path}")

        print(f"📄 Extracting: {docx_path.name}")

        # Load document
        doc = Document(docx_path)

        # Extract metadata từ docx properties
        metadata = self._extract_metadata(doc, docx_path)

        # Extract text và structure
        text, structure = self._extract_text_and_structure(doc)

        # Extract tables
        tables = self._extract_tables(doc)

        # Calculate statistics
        statistics = self._calculate_statistics(text, structure, tables)

        return ExtractedContent(
            text=text,
            metadata=metadata,
            structure=structure,
            tables=tables,
            statistics=statistics,
        )

    def _extract_metadata(self, doc: Document, file_path: Path) -> Dict:
        """Extract metadata từ document properties"""
        core_props = doc.core_properties

        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path.absolute()),
            "title": core_props.title or file_path.stem,
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": (
                core_props.created.isoformat() if core_props.created else None
            ),
            "modified": (
                core_props.modified.isoformat() if core_props.modified else None
            ),
        }

        # Try to extract document number/type từ title hoặc filename
        doc_info = self._parse_document_info(metadata["title"])
        metadata.update(doc_info)

        return metadata

    def _parse_document_info(self, title: str) -> Dict:
        """
        Parse thông tin văn bản từ title:
        - Luật số XX/YYYY-QH
        - Nghị định số XX/YYYY-NĐ-CP
        - Thông tư số XX/YYYY-TT-BTC
        """
        info = {"doc_type": "unknown", "doc_number": "", "doc_year": ""}

        # Luật
        law_match = re.search(r"Luật\s+(?:số\s+)?(\d+)/(\d{4})-?(\w+)?", title, re.I)
        if law_match:
            info["doc_type"] = "Luật"
            info["doc_number"] = f"{law_match.group(1)}/{law_match.group(2)}"
            if law_match.group(3):
                info["doc_number"] += f"-{law_match.group(3)}"
            info["doc_year"] = law_match.group(2)

        # Nghị định
        decree_match = re.search(
            r"Nghị định\s+(?:số\s+)?(\d+)/(\d{4})-?(\w+-?\w+)?", title, re.I
        )
        if decree_match:
            info["doc_type"] = "Nghị định"
            info["doc_number"] = f"{decree_match.group(1)}/{decree_match.group(2)}"
            if decree_match.group(3):
                info["doc_number"] += f"-{decree_match.group(3)}"
            info["doc_year"] = decree_match.group(2)

        # Thông tư
        circular_match = re.search(
            r"Thông tư\s+(?:số\s+)?(\d+)/(\d{4})-?(\w+-?\w+)?", title, re.I
        )
        if circular_match:
            info["doc_type"] = "Thông tư"
            info["doc_number"] = (
                f"{circular_match.group(1)}/{circular_match.group(2)}"
            )
            if circular_match.group(3):
                info["doc_number"] += f"-{circular_match.group(3)}"
            info["doc_year"] = circular_match.group(2)

        return info

    def _extract_text_and_structure(self, doc: Document) -> Tuple[str, List[Dict]]:
        """
        Extract text và detect hierarchical structure
        Returns: (full_text, structure_list)
        """
        full_text = []
        structure = []
        current_hierarchy = {
            "phan": None,
            "chuong": None,
            "muc": None,
            "dieu": None,
        }

        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Paragraph
                para = Paragraph(element, doc)
                text = para.text.strip()

                if not text:
                    continue

                # Detect structural elements
                struct_info = self._detect_structure(text, current_hierarchy)

                if struct_info:
                    structure.append(struct_info)
                    # Update hierarchy
                    for key in ["phan", "chuong", "muc", "dieu"]:
                        if key in struct_info:
                            current_hierarchy[key] = struct_info[key]

                full_text.append(text)

            elif isinstance(element, CT_Tbl):
                # Table - mark position
                full_text.append("[TABLE]")

        return "\n".join(full_text), structure

    def _detect_structure(
        self, text: str, current_hierarchy: Dict
    ) -> Optional[Dict]:
        """Detect loại cấu trúc của paragraph"""

        # Check Phần
        phan_match = re.match(self.patterns["phan"], text, re.I)
        if phan_match:
            return {
                "type": "phan",
                "phan": text,
                "text": text,
                "level": 1,
            }

        # Check Chương
        chuong_match = re.match(self.patterns["chuong"], text, re.I)
        if chuong_match:
            return {
                "type": "chuong",
                "chuong": text,
                "phan": current_hierarchy.get("phan"),
                "text": text,
                "level": 2,
            }

        # Check Mục
        muc_match = re.match(self.patterns["muc"], text, re.I)
        if muc_match:
            return {
                "type": "muc",
                "muc": text,
                "chuong": current_hierarchy.get("chuong"),
                "phan": current_hierarchy.get("phan"),
                "text": text,
                "level": 3,
            }

        # Check Điều
        dieu_match = re.match(self.patterns["dieu"], text, re.I)
        if dieu_match:
            dieu_num = dieu_match.group(1)
            # Extract title (text after "Điều X.")
            title = re.sub(r"^Điều\s+\d+[a-z]?[:\s.]+", "", text, flags=re.I).strip()

            return {
                "type": "dieu",
                "dieu": dieu_num,
                "dieu_text": text,
                "dieu_title": title,
                "muc": current_hierarchy.get("muc"),
                "chuong": current_hierarchy.get("chuong"),
                "phan": current_hierarchy.get("phan"),
                "text": text,
                "level": 4,
            }

        return None

    def _extract_tables(self, doc: Document) -> List[Dict]:
        """Extract tables từ document"""
        tables_data = []

        for idx, table in enumerate(doc.tables):
            table_data = {
                "table_id": idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": [],
            }

            # Extract table content
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["data"].append(row_data)

            tables_data.append(table_data)

        return tables_data

    def _calculate_statistics(
        self, text: str, structure: List[Dict], tables: List[Dict]
    ) -> Dict:
        """Calculate statistics về document"""

        stats = {
            "char_count": len(text),
            "line_count": len(text.split("\n")),
            "word_count": len(text.split()),
            "table_count": len(tables),
        }

        # Count structural elements
        for struct_type in ["phan", "chuong", "muc", "dieu"]:
            count = sum(1 for s in structure if s.get("type") == struct_type)
            stats[f"{struct_type}_count"] = count

        return stats

    def to_json(self, content: ExtractedContent) -> Dict:
        """Convert ExtractedContent sang JSON format"""
        return {
            "metadata": content.metadata,
            "content": {"full_text": content.text},
            "structure": content.structure,
            "tables": content.tables,
            "statistics": content.statistics,
        }

    def export_to_md(self, content: ExtractedContent, output_path: Path) -> None:
        """
        Export extracted content sang Markdown format
        với YAML frontmatter (compatible với md_processor)
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Build YAML frontmatter
        yaml_lines = ["---"]
        for key, value in content.metadata.items():
            if value:
                yaml_lines.append(f'{key}: "{value}"')
        yaml_lines.append("---")
        yaml_lines.append("")

        # Combine
        md_content = "\n".join(yaml_lines) + "\n" + content.text

        # Write
        output_path.write_text(md_content, encoding="utf-8")
        print(f"   💾 Exported to: {output_path}")


# ============ USAGE EXAMPLE ============

if __name__ == "__main__":
    from pathlib import Path

    # Example: Extract Luật Đấu thầu 2023
    docx_file = Path("data/raw/Luat chinh/Luat Dau thau 2023.docx")

    if docx_file.exists():
        extractor = DocxExtractor()
        content = extractor.extract(docx_file)

        print("\n📊 Extraction Statistics:")
        for key, value in content.statistics.items():
            print(f"   {key}: {value}")

        print(f"\n📝 Metadata:")
        for key, value in content.metadata.items():
            print(f"   {key}: {value}")

        print(f"\n🏗️ Structure Elements: {len(content.structure)}")

        # Export to JSON
        json_output = extractor.to_json(content)

        # Export to MD
        md_output = Path("data/processed/Luat-Dau-thau-2023.md")
        extractor.export_to_md(content, md_output)

        print("\n✅ Extraction complete!")
    else:
        print(f"❌ File not found: {docx_file}")
